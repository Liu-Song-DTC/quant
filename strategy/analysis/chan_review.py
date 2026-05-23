#!/usr/bin/env python3
"""
量化系统每日复盘 & 自我进化引擎

设计理念（融合盘后复盘方法论）:
  复盘 ≠ 单纯看报告。复盘 = 体检诊断 → 发现偏差 → 驱动系统参数自调整。

复盘6步法（量化落地版）:
  Step 1: 大盘环境深度体检 — 指数状态、情绪周期、量价关系、风格偏向
  Step 2: 板块与主线逻辑拆解 — 行业趋势浓度、梯队完整度、资金流向
  Step 3: 涨停板与情绪监控 — 涨跌停统计、连板高度、赚钱效应
  Step 4: 持仓股与自选股排雷 — 关键支撑/均线破位检查、逻辑验证
  Step 5: 缠论结构全面扫描 — 走势类型、买卖点、背驰、多级别确认
  Step 6: 生成作战计划 + 系统自进化建议

使用方式:
  cd strategy && python analysis/chan_review.py                      # 复盘最新交易日
  cd strategy && python analysis/chan_review.py --date 2026-05-15    # 复盘指定日期
  cd strategy && python analysis/chan_review.py --top 20 --evolve    # 复盘+自动进化
"""

import os
import sys
import argparse
import json
import warnings
from pathlib import Path
from datetime import datetime, timedelta
from dataclasses import dataclass, field
from typing import Dict, List, Optional, Tuple
from collections import defaultdict

import numpy as np
import pandas as pd

warnings.filterwarnings('ignore')

SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
STRATEGY_DIR = os.path.dirname(SCRIPT_DIR)
PROJECT_ROOT = os.path.dirname(STRATEGY_DIR)
sys.path.insert(0, STRATEGY_DIR)

from core.chan_theory import (
    compute_enhanced_chan_output,
    analyze_bottom_fractal,
    detect_stroke_mmd,
    check_bi_trend_depletion,
    get_structure_stop_price,
)
from core.market_regime_detector import MarketRegimeDetector
from core.industry_mapping import INDUSTRY_KEYWORDS, get_industry_category
from core.chan_theory import _calc_ema

# ==================== 配置 ====================

DATA_DIR = os.path.join(PROJECT_ROOT, 'data', 'stock_data', 'backtrader_data')
INDEX_FILE = os.path.join(DATA_DIR, 'sh000001_qfq.csv')
CONFIG_PATH = os.path.join(STRATEGY_DIR, 'config', 'factor_config.yaml')
REVIEW_OUTPUT_DIR = os.path.join(STRATEGY_DIR, 'daily_review')
EVOLUTION_LOG = os.path.join(REVIEW_OUTPUT_DIR, 'evolution_history.json')

DEFAULT_MIN_VOLUME = 20_000_000   # 最小成交额2000万
DEFAULT_MIN_TURNOVER = 15_000_000  # 备用：成交额门槛

# 关键均线（排雷用）
KEY_MA_PERIODS = [5, 10, 20, 60]

# 情绪周期阈值
SENTIMENT_THRESHOLDS = {
    'ice': {'limit_up_max': 30, 'adv_pct_max': 0.25},        # 冰点: 涨停<30, 上涨占比<25%
    'repair': {'limit_up_range': (30, 60), 'adv_pct_range': (0.25, 0.45)},  # 修复
    'climax': {'limit_up_min': 60, 'adv_pct_min': 0.45},     # 高潮: 涨停>60, 上涨>45%
    'divergence': {'adv_pct_range': (0.30, 0.55)},            # 分歧: 涨跌互现
    'retreat': {'limit_up_max': 40, 'adv_pct_max': 0.35},     # 退潮: 涨停减少
}


# ==================== 数据结构 ====================

@dataclass
class MarketVitalSigns:
    """大盘体检报告"""
    date: str = ''

    # 指数状态
    index_name: str = '上证指数'
    index_close: float = 0.0
    index_change_pct: float = 0.0
    index_volume: str = ''  # 放量/缩量/平量
    volume_ratio: float = 1.0

    # 市场状态
    regime: str = ''          # 牛市/震荡/熊市
    regime_confidence: float = 0.0
    sentiment_cycle: str = ''  # 冰点/修复/高潮/分歧/退潮
    volatility_pct: float = 0.0
    momentum_score: float = 0.0

    # 赚钱效应
    advancing: int = 0
    declining: int = 0
    up_pct: float = 0.0
    limit_up_count: int = 0
    limit_down_count: int = 0
    limit_up_ratio: float = 0.0  # 涨停/跌停比

    # 风格偏向
    style_regime: str = ''
    bear_risk: bool = False

    # 指数共振
    index_divergence: bool = False  # 上证/深成/创业板是否分化

    # 综合判定
    summary: str = ''


@dataclass
class IndustryMainLine:
    """行业主线分析"""
    industry: str = ''
    stock_count: int = 0

    # 趋势分布
    up_trend_pct: float = 0.0
    down_trend_pct: float = 0.0
    consolidation_pct: float = 0.0

    # 缠论买点浓度
    b1_count: int = 0
    b2_count: int = 0
    b3_count: int = 0
    total_buy_signals: int = 0
    total_sell_signals: int = 0

    # 资金流向
    avg_volume_ratio: float = 1.0
    net_flow_score: float = 0.0

    # 梯队完整度 (量化近似)
    has_leader: bool = False      # 有龙头（涨>5%+放量>2倍均量）
    has_mid_cap: bool = False     # 有中军（涨>2%+成交额>行业平均）
    has_follower: bool = False    # 有跟风（行业内涨跌比>50%）
    formation_score: float = 0.0  # 梯队完整度 0-1

    # 综合评分
    hot_rank: int = 0
    avg_composite_score: float = 0.0
    summary: str = ''


@dataclass
class StockDiagnosis:
    """个股诊断（缠论 + 排雷）"""
    symbol: str
    name: str = ''
    industry: str = ''

    # 基础数据
    close: float = 0.0
    change_pct: float = 0.0
    volume_ratio: float = 1.0

    # 缠论结构
    trend_type: int = 0          # 2=上涨趋势, 1=盘整, -2=下跌趋势
    trend_strength: float = 0.0
    pivot_position: int = 0      # -1=中枢下, 0=中枢内, 1=中枢上
    pivot_zg: float = np.nan
    pivot_zd: float = np.nan
    stroke_direction: int = 0    # 当前笔方向
    segment_direction: int = 0   # 当前线段方向

    # 买卖点
    buy_point: int = 0
    sell_point: int = 0
    buy_confidence: float = 0.0
    sell_confidence: float = 0.0
    second_buy: bool = False
    second_buy_conf: float = 0.0
    confirmed_buy: bool = False
    confirmed_sell: bool = False
    signal_level: int = 0
    buy_strength: float = 0.0
    sell_strength: float = 0.0

    # 背驰
    bottom_divergence: bool = False
    top_divergence: bool = False

    # 排雷检查
    ma_breakdowns: List[str] = field(default_factory=list)   # 均线破位列表
    pivot_breakdown: bool = False   # 中枢下方运行
    volume_anomaly: str = ''        # 倍量/堆量/缩量
    bi_trend_depletion: bool = False

    # 均线排列
    alignment_score: float = 0.0

    # 综合评分 & 等级
    composite_score: float = 0.0
    opportunity_rank: str = ''     # A/B/C/D
    risk_level: str = ''           # 高危/注意/安全

    # 止损参考
    stop_loss_price: float = 0.0


@dataclass
class LimitUpDiagnosis:
    """涨停板缠论归因分析"""
    symbol: str
    name: str = ''
    industry: str = ''
    close: float = 0.0
    change_pct: float = 0.0
    volume_ratio: float = 1.0

    # 涨停前5日的缠论结构快照
    pre_trend_type: int = 0         # 涨停前的走势类型
    pre_pivot_position: int = 0     # 涨停前相对中枢位置
    pre_stroke_direction: int = 0   # 涨停前笔方向
    pre_buy_point: int = 0          # 涨停前是否有买点
    pre_buy_confidence: float = 0.0
    pre_bottom_divergence: bool = False  # 涨停前底背驰
    pre_alignment_score: float = 0.0
    pre_consolidation_days: int = 0      # 涨停前盘整天数（近似）

    # 涨停日的信号
    cur_buy_point: int = 0
    cur_buy_confidence: float = 0.0
    cur_bottom_divergence: bool = False

    # 结构特征（涨停日附近的）
    pivot_breakout: bool = False     # 是否突破中枢上沿
    pivot_zg: float = np.nan
    pivot_zd: float = np.nan
    above_pivot_pct: float = 0.0    # 收盘价超过ZG的百分比
    below_pivot_pct: float = 0.0    # 开盘前在ZD下方的百分比

    # 量能特征
    vol_spike: float = 1.0          # 涨停日相对5日均量
    bottom_fx_quality: float = 0.0  # 最近底分型质量

    # 涨停归因分类
    limit_up_type: str = ''         # 中枢突破板/底背驰反转板/二买确认板/趋势加速板/笔耗尽反转板/盘整突破板
    type_confidence: float = 0.0    # 归因置信度
    chan_explanation: str = ''      # 缠论解释


@dataclass
class LimitUpSummary:
    """涨停板共性总结"""
    date: str
    total_limit_up: int
    limit_up_stocks: List[LimitUpDiagnosis] = field(default_factory=list)

    # 按类型分布
    type_distribution: Dict[str, int] = field(default_factory=dict)
    type_avg_confidence: Dict[str, float] = field(default_factory=dict)

    # 共性特征
    common_traits: List[str] = field(default_factory=list)      # 共性发现
    industry_concentration: Dict[str, int] = field(default_factory=dict)  # 涨停行业集中度
    dominant_pattern: str = ''      # 今日主导涨停模式


@dataclass
class EvolutionPatch:
    """系统自进化补丁"""
    section: str          # 配置section
    key: str              # 配置key
    current_value: float
    suggested_value: float
    reason: str
    confidence: float     # 建议置信度 0-1
    urgency: str          # urgent/high/medium/low


@dataclass
class DailyReviewReport:
    """完整的每日复盘报告"""
    date: str
    market: MarketVitalSigns
    main_lines: List[IndustryMainLine]
    opportunities: Dict[str, List[StockDiagnosis]]
    risk_alerts: Dict[str, List[StockDiagnosis]]
    portfolio_suggestions: List[StockDiagnosis]
    limit_up_summary: Optional[LimitUpSummary] = None
    evolution_patches: List[EvolutionPatch] = field(default_factory=list)
    all_diagnoses: List[StockDiagnosis] = field(default_factory=list)  # 全量个股诊断
    pred_accuracy: Dict[str, any] = field(default_factory=dict)       # 预测准确性评估
    portfolio_holdings: List[Dict] = field(default_factory=list)      # 持仓评估
    today_selections: List[Dict] = field(default_factory=list)        # 当日选股结果
    replacement_plan: List[Dict] = field(default_factory=list)        # 持仓替换计划
    portfolio_cash: float = 0.0                                        # 可用资金
    accuracy_history: List[Dict] = field(default_factory=list)         # 历史准确率
    summary: Dict[str, any] = field(default_factory=dict)


# ==================== 复盘引擎 ====================

class QuantReviewEngine:
    """量化复盘引擎 — 体检诊断 → 机会发现 → 系统进化"""

    def __init__(self, data_dir: str = None, min_volume: float = None):
        self.data_dir = data_dir or DATA_DIR
        self.min_volume = min_volume or DEFAULT_MIN_VOLUME
        self.output_dir = REVIEW_OUTPUT_DIR
        os.makedirs(self.output_dir, exist_ok=True)

        self._price_cache = {}
        self._industry_map = self._load_industry_map()

    def _load_industry_map(self) -> dict:
        """从基本面数据加载 code -> 所处行业 映射表"""
        import glob
        industry_map = {}
        fund_dir = os.path.join(PROJECT_ROOT, 'data', 'stock_data', 'fundamental_data')
        if not os.path.exists(fund_dir):
            return industry_map
        try:
            for fp in glob.glob(os.path.join(fund_dir, '*.csv')):
                try:
                    code = os.path.basename(fp).replace('.csv', '').zfill(6)
                    df = pd.read_csv(fp, nrows=1)
                    if '所处行业' in df.columns:
                        raw = str(df['所处行业'].iloc[0])
                        if raw and raw != 'nan':
                            cleaned = raw.replace('Ⅱ', '').replace('Ⅲ', '').replace('Ⅳ', '').strip()
                            industry_map[code] = cleaned
                except Exception:
                    pass
        except Exception:
            pass
        return industry_map

    # ========== 数据层 ==========

    def _find_latest_date(self) -> str:
        import glob
        files = glob.glob(os.path.join(self.data_dir, '*_qfq.csv'))
        if not files:
            return datetime.now().strftime('%Y-%m-%d')
        for f in files:
            try:
                df = pd.read_csv(f, nrows=5)
                for col in ['datetime', 'date']:
                    if col in df.columns and len(df) > 0:
                        return str(pd.to_datetime(df[col].max()).date())
            except Exception:
                continue
        return datetime.now().strftime('%Y-%m-%d')

    def _load_index_data(self, date: str, lookback: int = 250) -> pd.DataFrame:
        if not os.path.exists(INDEX_FILE):
            return pd.DataFrame()
        df = pd.read_csv(INDEX_FILE)
        date_col = 'datetime' if 'datetime' in df.columns else 'date'
        if date_col not in df.columns:
            return pd.DataFrame()
        if date_col != 'date':
            df = df.rename(columns={date_col: 'date'})
        df['date'] = df['date'].astype(str)
        return df[df['date'] <= date].tail(lookback).copy()

    def _load_stock_data(self, symbol: str, date: str, lookback: int = 250) -> Optional[pd.DataFrame]:
        cache_key = f'{symbol}_{date}_{lookback}'
        if cache_key in self._price_cache:
            return self._price_cache[cache_key]

        fp = os.path.join(self.data_dir, f'{symbol}_qfq.csv')
        if not os.path.exists(fp):
            return None
        try:
            df = pd.read_csv(fp)
            # 统一日期列名
            date_col = 'datetime' if 'datetime' in df.columns else 'date'
            if date_col != 'date':
                df = df.rename(columns={date_col: 'date'})
            if 'date' not in df.columns or len(df) < 60:
                return None
            df['date'] = df['date'].astype(str)
            df = df[df['date'] <= date].tail(lookback).copy()
            if len(df) < 60:
                return None

            # 标准化列名
            col_map = {}
            for std_col in ['open', 'high', 'low', 'close', 'volume']:
                for c in df.columns:
                    if c.lower() == std_col:
                        col_map[std_col] = c
                        break
            df = df.rename(columns={v: k for k, v in col_map.items()})

            # 成交额列（如果有）
            for c in ['turnover', 'amount', '成交额']:
                if c in df.columns:
                    col_map['turnover'] = c
                    break

            self._price_cache[cache_key] = df
            return df
        except Exception:
            return None

    def _list_available_stocks(self) -> List[str]:
        import glob
        files = glob.glob(os.path.join(self.data_dir, '*_qfq.csv'))
        skip_prefixes = ('sh000', 'sh399', 'sz399', 'sh688', 'bj')
        symbols = []
        for f in files:
            bn = os.path.basename(f)
            sym = bn.replace('_qfq.csv', '')
            # 跳过指数
            if any(sym.startswith(p) for p in skip_prefixes):
                continue
            # 只保留纯数字代码（6位）+ 可选sh/sz前缀
            code = sym[2:] if sym[:2] in ('sh', 'sz') else sym
            if code.isdigit() and len(code) == 6:
                symbols.append(sym)
        return sorted(symbols)

    # 股票名称缓存（类级别，避免重复加载CSV）
    _stock_name_cache: Dict[str, str] = {}

    def _get_stock_name(self, symbol: str) -> str:
        """获取股票名称，使用CSV文件并缓存"""
        if symbol in self._stock_name_cache:
            return self._stock_name_cache[symbol]

        # 尝试多个可能的文件名
        for fname in ['stock_list.csv', 'stock_list_full.csv']:
            metadata_file = os.path.join(
                PROJECT_ROOT, 'data', 'stock_data', 'stock_metadata', fname
            )
            if os.path.exists(metadata_file):
                try:
                    df = pd.read_csv(metadata_file, dtype={'symbol': str})
                    if 'symbol' in df.columns and 'name' in df.columns:
                        name_map = dict(zip(df['symbol'].astype(str).str.strip(),
                                           df['name'].astype(str).str.strip()))
                        # 缓存全部映射
                        self._stock_name_cache.update(name_map)
                        return self._stock_name_cache.get(symbol, '')
                except Exception:
                    pass
        return ''

    # ========== Step 1: 大盘环境深度体检 ==========

    def examine_market_vitals(self, date: str, stock_list: List[str]) -> MarketVitalSigns:
        """大盘环境深度体检"""
        v = MarketVitalSigns(date=date)

        index_df = self._load_index_data(date)
        if len(index_df) < 60:
            v.summary = '数据不足，无法完成大盘体检'
            return v

        close = index_df['close'].values.astype(np.float64)
        v.index_close = float(close[-1])
        v.index_change_pct = float((close[-1] / close[-2] - 1) * 100) if len(close) >= 2 else 0.0

        # 成交量判断
        vol = index_df['volume'].values.astype(np.float64) if 'volume' in index_df.columns else np.ones(len(close))
        vol_ma20 = np.mean(vol[-21:-1]) if len(vol) >= 21 else vol[-1]
        v.volume_ratio = float(vol[-1] / vol_ma20) if vol_ma20 > 0 else 1.0
        if v.volume_ratio >= 1.3:
            v.index_volume = '放量'
        elif v.volume_ratio <= 0.7:
            v.index_volume = '缩量'
        else:
            v.index_volume = '平量'

        # 市场状态检测
        detector = MarketRegimeDetector()
        regime_df = detector.generate(index_df)
        last_row = regime_df.iloc[-1]
        regime_labels = {1: '牛市', 0: '震荡', -1: '熊市'}
        v.regime = regime_labels.get(int(last_row.get('regime', 0)), '未知')
        v.regime_confidence = float(last_row.get('confidence', 0))
        v.momentum_score = float(last_row.get('momentum_score', 0))
        v.volatility_pct = float(last_row.get('volatility', 0))
        v.style_regime = str(last_row.get('style_regime', ''))
        v.bear_risk = bool(last_row.get('bear_risk', False))

        # 赚钱效应统计（抽样300只）
        adv, dec, limit_up, limit_down = 0, 0, 0, 0
        for sym in stock_list[:300]:
            df = self._load_stock_data(sym, date, lookback=5)
            if df is None or len(df) < 2:
                continue
            chg = (df['close'].iloc[-1] / df['close'].iloc[-2] - 1)
            if chg > 0:
                adv += 1
            elif chg < 0:
                dec += 1
            if chg >= 0.098:
                limit_up += 1
            elif chg <= -0.098:
                limit_down += 1

        v.advancing = adv
        v.declining = dec
        v.up_pct = adv / max(adv + dec, 1) * 100
        v.limit_up_count = limit_up
        v.limit_down_count = limit_down
        v.limit_up_ratio = limit_up / max(limit_down, 1)

        # 情绪周期判断
        v.sentiment_cycle = self._classify_sentiment(limit_up, limit_down, v.up_pct / 100)

        # 指数共振（简化版：检查市场宽度）
        v.index_divergence = (v.up_pct < 45 and v.index_change_pct > 0.3) or \
                             (v.up_pct > 55 and v.index_change_pct < -0.3)

        # 综合判定
        parts = []
        parts.append(f"{v.regime}{'放量' if v.index_volume == '放量' else ''}")
        parts.append(f"情绪: {v.sentiment_cycle}")
        conditions = []
        if v.bear_risk:
            conditions.append("⚠熊市风险")
        if v.index_divergence:
            conditions.append("指数分化")
        if v.limit_up_ratio >= 3:
            conditions.append("赚钱效应强")
        elif v.limit_up_ratio <= 0.5:
            conditions.append("亏钱效应强")
        parts.append(' | '.join(conditions) if conditions else '正常')
        v.summary = '，'.join(parts)

        return v

    def _classify_sentiment(self, limit_up: int, limit_down: int, up_pct: float) -> str:
        """情绪周期分类: 冰点→修复→高潮→分歧→退潮"""
        if limit_up < 30 and up_pct < 0.25:
            return '冰点'
        if limit_up > 60 and up_pct > 0.50:
            return '高潮'
        if limit_down > limit_up * 0.6:
            return '退潮'
        if 0.30 <= up_pct <= 0.55 and abs(limit_up - limit_down) <= limit_up * 0.4:
            return '分歧'
        if 30 <= limit_up <= 60 and 0.25 <= up_pct <= 0.45:
            return '修复'
        return '分歧'

    # ========== Step 2: 板块与主线逻辑拆解 ==========

    def analyze_industry_main_lines(
        self, diagnoses: List[StockDiagnosis]
    ) -> List[IndustryMainLine]:
        """行业主线分析 — 包含梯队完整度"""
        ind_data = defaultdict(lambda: {
            'stocks': [], 'buys': 0, 'sells': 0,
            'b1': 0, 'b2': 0, 'b3': 0,
            'up_trend': 0, 'down_trend': 0, 'cons': 0,
            'vol_ratios': [], 'composites': [],
            'leaders': 0, 'mid_caps': 0, 'followers': 0,
        })

        for d in diagnoses:
            ind = d.industry or '其他'
            dd = ind_data[ind]
            dd['stocks'].append(d)
            if d.buy_point > 0:
                dd['buys'] += 1
            if d.sell_point > 0:
                dd['sells'] += 1
            if d.buy_point == 1:
                dd['b1'] += 1
            elif d.buy_point == 2:
                dd['b2'] += 1
            elif d.buy_point == 3:
                dd['b3'] += 1

            if d.trend_type == 2:
                dd['up_trend'] += 1
            elif d.trend_type == -2:
                dd['down_trend'] += 1
            else:
                dd['cons'] += 1

            dd['vol_ratios'].append(d.volume_ratio)
            dd['composites'].append(d.composite_score)

            # 梯队检测
            if d.change_pct > 5 and d.volume_ratio > 2:
                dd['leaders'] += 1
            if d.change_pct > 2 and d.volume_ratio > 1.5:
                dd['mid_caps'] += 1
            if d.change_pct > 0:
                dd['followers'] += 1

        results = []
        for ind, dd in ind_data.items():
            n = len(dd['stocks'])
            if n < 3:
                continue

            ml = IndustryMainLine(industry=ind, stock_count=n)
            ml.up_trend_pct = dd['up_trend'] / n * 100
            ml.down_trend_pct = dd['down_trend'] / n * 100
            ml.consolidation_pct = dd['cons'] / n * 100
            ml.b1_count = dd['b1']
            ml.b2_count = dd['b2']
            ml.b3_count = dd['b3']
            ml.total_buy_signals = dd['buys']
            ml.total_sell_signals = dd['sells']
            ml.avg_volume_ratio = np.mean(dd['vol_ratios']) if dd['vol_ratios'] else 1.0
            ml.avg_composite_score = np.mean(dd['composites']) if dd['composites'] else 0.0

            # 梯队完整度
            ml.has_leader = dd['leaders'] > 0
            ml.has_mid_cap = dd['mid_caps'] > 1
            ml.has_follower = (dd['followers'] / n) > 0.5 if n > 0 else False
            ml.formation_score = (
                (0.4 if ml.has_leader else 0) +
                (0.3 if ml.has_mid_cap else 0) +
                (0.3 if ml.has_follower else 0)
            )
            ml.net_flow_score = dd['buys'] - dd['sells']

            # 综合判定
            if ml.formation_score >= 0.7 and ml.total_buy_signals >= 3:
                ml.summary = '★ 强主线，梯队完整，可积极操作'
            elif ml.formation_score >= 0.4 and ml.total_buy_signals >= 1:
                ml.summary = '○ 次主线，有结构但梯队不全，谨慎参与'
            elif ml.total_sell_signals > ml.total_buy_signals:
                ml.summary = '✗ 资金出逃，回避'
            else:
                ml.summary = '— 无明显方向，观望'

            results.append(ml)

        results.sort(key=lambda x: (x.formation_score + x.avg_composite_score / 3), reverse=True)
        for i, r in enumerate(results):
            r.hot_rank = i + 1
        return results

    # ========== Step 3+5: 涨停监控 & 缠论扫描（合并） ==========

    def diagnose_stock(
        self, symbol: str, date: str, index_df: pd.DataFrame = None
    ) -> Optional[StockDiagnosis]:
        """个股全面诊断 — 缠论结构 + 排雷检查"""
        df = self._load_stock_data(symbol, date)
        if df is None or len(df) < 60:
            return None

        # 流动性过滤 — 优先使用amount列(成交额)，其次 volume*close
        if 'amount' in df.columns:
            avg_turnover = df['amount'].tail(20).mean()
        elif 'volume' in df.columns and 'close' in df.columns:
            avg_turnover = (df['volume'].tail(20) * df['close'].tail(20)).mean()
        else:
            avg_turnover = float('inf')
        if avg_turnover < self.min_volume:
                return None

        n = len(df)
        close = df['close'].values.astype(np.float64)
        high = df['high'].values.astype(np.float64)
        low = df['low'].values.astype(np.float64)
        volume = df['volume'].values.astype(np.float64) if 'volume' in df.columns else np.ones(n)

        # === 缠论增强输出 ===
        ema12 = _calc_ema(close, 12)
        ema26 = _calc_ema(close, 26)
        macd_line = ema12 - ema26
        macd_hist = 2 * (macd_line - _calc_ema(macd_line, 9))
        ema20 = _calc_ema(close, 20)
        ema60 = _calc_ema(close, 60)
        ema120 = _calc_ema(close, 120)

        chan = compute_enhanced_chan_output(
            close, high, low, ema20, ema60, ema120, macd_hist, volume,
        )

        idx = -1
        d = StockDiagnosis(symbol=symbol)
        d.name = self._get_stock_name(symbol)
        try:
            raw_ind = self._industry_map.get(d.symbol, '')
            d.industry = get_industry_category(raw_ind) if raw_ind else '未归类'
        except Exception:
            raw_ind = self._industry_map.get(d.symbol, '')
            d.industry = get_industry_category(raw_ind) if raw_ind else '未归类'

        d.close = float(close[idx])
        d.change_pct = float((close[idx] / close[idx - 1] - 1) * 100) if idx - 1 >= -n else 0.0
        d.volume_ratio = float(volume[idx] / np.mean(volume[-21:-1])) if len(volume) >= 21 and np.mean(volume[-21:-1]) > 0 else 1.0

        # 缠论结构
        d.trend_type = int(chan['trend_type'][idx])
        d.trend_strength = float(chan['trend_strength'][idx])
        d.pivot_position = int(chan['pivot_position'][idx])
        d.pivot_zg = float(chan['pivot_zg'][idx]) if not np.isnan(chan['pivot_zg'][idx]) else np.nan
        d.pivot_zd = float(chan['pivot_zd'][idx]) if not np.isnan(chan['pivot_zd'][idx]) else np.nan
        d.stroke_direction = int(chan['stroke_direction'][idx])
        d.segment_direction = int(chan['segment_direction'][idx])

        # 买卖点
        d.buy_point = int(chan['buy_point'][idx])
        d.sell_point = int(chan['sell_point'][idx])
        d.buy_confidence = float(chan['buy_confidence'][idx])
        d.sell_confidence = float(chan['sell_confidence'][idx])
        d.second_buy = bool(chan.get('second_buy_point', np.zeros(n))[idx])
        d.second_buy_conf = float(chan.get('second_buy_confidence', np.zeros(n))[idx])
        d.confirmed_buy = bool(chan['confirmed_buy'][idx])
        d.confirmed_sell = bool(chan['confirmed_sell'][idx])
        d.signal_level = int(chan['signal_level'][idx])
        d.buy_strength = float(chan['buy_strength'][idx])
        d.sell_strength = float(chan['sell_strength'][idx])

        # 背驰
        d.bottom_divergence = bool(chan.get('hidden_bottom_divergence', np.zeros(n))[idx])
        d.top_divergence = bool(chan.get('hidden_top_divergence', np.zeros(n))[idx])

        # 均线排列
        d.alignment_score = float(chan['alignment_score'][idx])

        # 笔趋势耗尽
        d.bi_trend_depletion = bool(chan.get('bi_td', np.zeros(n))[idx])

        # === 排雷检查 ===
        # 1. 均线破位
        d.ma_breakdowns = []
        for ma_period in KEY_MA_PERIODS:
            ma = _calc_ema(close, ma_period)
            if not np.isnan(ma[idx]) and close[idx] < ma[idx] * 0.98:
                d.ma_breakdowns.append(f'跌破MA{ma_period}')

        # 2. 中枢破位
        if d.pivot_position == -1 and d.trend_type <= 0:
            d.pivot_breakdown = True

        # 3. 量异常
        if d.volume_ratio >= 2.5:
            d.volume_anomaly = '倍量'
        elif d.volume_ratio >= 1.5:
            d.volume_anomaly = '放量'
        elif d.volume_ratio <= 0.4:
            d.volume_anomaly = '缩量'
        else:
            d.volume_anomaly = '正常'

        # === 综合评分 ===
        d.composite_score = self._compute_composite(d)
        d.opportunity_rank = self._rank_opportunity(d)
        d.risk_level = self._assess_risk(d)
        d.stop_loss_price = float(chan.get('structure_stop_price', np.full(n, np.nan))[idx])
        if np.isnan(d.stop_loss_price):
            d.stop_loss_price = d.close * 0.93

        return d

    def _compute_composite(self, d: StockDiagnosis) -> float:
        """综合评分 [-1, +1]"""
        score = 0.0
        # 趋势 (25%)
        if d.trend_type == 2:
            score += 0.25
        elif d.trend_type == -2:
            score -= 0.25

        # 买卖点 (35%)
        if d.confirmed_buy:
            score += 0.35 * d.buy_strength
        elif d.confirmed_sell:
            score -= 0.35 * d.sell_strength
        elif d.buy_point > 0:
            score += 0.20 * d.buy_confidence
        elif d.sell_point > 0:
            score -= 0.20 * d.sell_confidence

        # 二买 (10%)
        if d.second_buy:
            score += 0.10 * d.second_buy_conf

        # 中枢位置 (10%)
        if d.pivot_position == 1:
            score += 0.08
        elif d.pivot_position == -1 and d.trend_type >= 0:
            score += 0.05

        # 均线对齐 (5%)
        score += d.alignment_score * 0.05

        # 排雷惩罚 (10%)
        if d.ma_breakdowns:
            score -= len(d.ma_breakdowns) * 0.05
        if d.pivot_breakdown:
            score -= 0.05
        if d.top_divergence:
            score -= 0.10

        # 笔趋势耗尽 (5%)
        if d.bi_trend_depletion and d.stroke_direction == -1:
            score += 0.05

        return float(np.clip(score, -1.0, 1.0))

    def _rank_opportunity(self, d: StockDiagnosis) -> str:
        """机会等级 A/B/C/D"""
        if d.confirmed_buy and d.signal_level >= 3 and d.buy_confidence >= 0.6 and d.composite_score >= 0.5:
            return 'A'
        if d.second_buy and d.second_buy_conf >= 0.6 and d.composite_score >= 0.4:
            return 'A'
        if d.buy_point >= 2 and d.buy_confidence >= 0.5 and d.composite_score >= 0.3:
            return 'B'
        if d.buy_point == 1 and d.composite_score >= 0.1:
            return 'C'
        if d.composite_score <= -0.5:
            return 'D-'
        return 'D'

    def _assess_risk(self, d: StockDiagnosis) -> str:
        """风险评估"""
        if d.confirmed_sell or d.sell_point >= 1:
            return '高危'
        if len(d.ma_breakdowns) >= 2 or d.pivot_breakdown:
            return '高危'
        if d.top_divergence or (d.bi_trend_depletion and d.stroke_direction == 1):
            return '注意'
        if d.trend_type == -2 and d.pivot_position <= 0:
            return '注意'
        return '安全'

    # ========== 机会 & 风险分类 ==========

    def classify_opportunities(
        self, diagnoses: List[StockDiagnosis],
        today_selections: List[Dict] = None,
    ) -> Dict[str, List[StockDiagnosis]]:
        """分类排序机会 — 优先使用策略选股结果，交叉缠论诊断做深度分析"""
        ops = {
            'A_strong_buy': [], 'B2_confirmed': [], 'B1_reversal': [],
            'B3_acceleration': [], 'double_confirmed': [], 'bottom_divergence': [],
        }

        if today_selections:
            # 使用策略选股结果，交叉缠论诊断
            diag_map = {d.symbol: d for d in diagnoses}
            for s in today_selections:
                code = str(s.get('股票代码', ''))
                score = float(s.get('综合评分', 0))
                d = diag_map.get(code)
                if d:
                    # 用策略评分替代纯缠论评分
                    d.composite_score = score
                    d.opportunity_rank = self._score_to_rank(score)
                else:
                    # 策略选了但缠论未覆盖，创建基础诊断
                    d = StockDiagnosis(
                        symbol=code,
                        name=str(s.get('股票名称', '')),
                        industry=str(s.get('行业', '')),
                        close=float(s.get('当前价格', 0)),
                        composite_score=score,
                        opportunity_rank=self._score_to_rank(score),
                    )
                ops = self._classify_one(d, ops)
        else:
            # 无策略选股时回退全市场缠论扫描
            for d in diagnoses:
                ops = self._classify_one(d, ops)

        ops['A_strong_buy'].sort(key=lambda x: x.composite_score, reverse=True)
        ops['B2_confirmed'].sort(key=lambda x: x.second_buy_conf, reverse=True)
        ops['B1_reversal'].sort(key=lambda x: x.buy_confidence, reverse=True)
        ops['B3_acceleration'].sort(key=lambda x: x.buy_confidence, reverse=True)
        ops['double_confirmed'].sort(key=lambda x: x.buy_strength, reverse=True)
        ops['bottom_divergence'].sort(key=lambda x: x.buy_confidence, reverse=True)
        return ops

    @staticmethod
    def _classify_one(d, ops):
        """将单只股票诊断归类到机会分类"""
        if d.opportunity_rank == 'A':
            ops['A_strong_buy'].append(d)
        if d.second_buy and d.second_buy_conf >= 0.45:
            ops['B2_confirmed'].append(d)
        if d.buy_point == 1 and d.buy_confidence >= 0.4:
            ops['B1_reversal'].append(d)
        if d.buy_point == 3:
            ops['B3_acceleration'].append(d)
        if d.confirmed_buy and d.signal_level >= 3:
            ops['double_confirmed'].append(d)
        if d.bottom_divergence and d.buy_point > 0:
            ops['bottom_divergence'].append(d)
        return ops

    def classify_risks(
        self, diagnoses: List[StockDiagnosis]
    ) -> Dict[str, List[StockDiagnosis]]:
        """分类排序风险"""
        risks = {
            'S1_top_reversal': [], 'S3_breakdown': [],
            'confirmed_sell': [], 'top_divergence': [],
            'ma_breakdown': [], 'trend_depletion': [],
        }
        for d in diagnoses:
            if d.sell_point == 1:
                risks['S1_top_reversal'].append(d)
            if d.sell_point == 3:
                risks['S3_breakdown'].append(d)
            if d.confirmed_sell:
                risks['confirmed_sell'].append(d)
            if d.top_divergence:
                risks['top_divergence'].append(d)
            if len(d.ma_breakdowns) >= 2:
                risks['ma_breakdown'].append(d)
            if d.bi_trend_depletion and d.stroke_direction == 1:
                risks['trend_depletion'].append(d)

        for k in risks:
            risks[k].sort(key=lambda x: x.sell_confidence if hasattr(x, 'sell_confidence') else 0, reverse=True)
        return risks

    # ========== Step 3 增强: 涨停板缠论归因分析 ==========

    LIMIT_UP_THRESHOLD = 9.5   # 涨停阈值(%), A股主板10%, 科创/创业20%

    def analyze_limit_up_stocks(
        self, date: str, stock_list: List[str], index_df: pd.DataFrame
    ) -> LimitUpSummary:
        """涨停板缠论归因分析 — 用缠论解释每个涨停的成因，总结共性"""

        summary = LimitUpSummary(date=date, total_limit_up=0)
        limit_up_diagnoses = []

        for sym in stock_list:
            d = self._diagnose_single_limit_up(sym, date, index_df)
            if d is not None:
                limit_up_diagnoses.append(d)

        summary.total_limit_up = len(limit_up_diagnoses)
        summary.limit_up_stocks = limit_up_diagnoses

        if not limit_up_diagnoses:
            return summary

        # 类型分布
        type_dist = defaultdict(int)
        type_confs = defaultdict(list)
        for d in limit_up_diagnoses:
            t = d.limit_up_type or '未分类'
            type_dist[t] += 1
            type_confs[t].append(d.type_confidence)

        summary.type_distribution = dict(type_dist)
        summary.type_avg_confidence = {
            k: np.mean(v) for k, v in type_confs.items()
        }

        # 行业集中度
        ind_counts = defaultdict(int)
        for d in limit_up_diagnoses:
            ind_counts[d.industry] += 1
        summary.industry_concentration = dict(
            sorted(ind_counts.items(), key=lambda x: x[1], reverse=True)[:10]
        )

        # 主导模式
        if type_dist:
            summary.dominant_pattern = max(type_dist, key=type_dist.get)

        # 共性发现
        summary.common_traits = self._extract_limit_up_commonalities(limit_up_diagnoses)

        return summary

    def _diagnose_single_limit_up(
        self, symbol: str, date: str, index_df: pd.DataFrame
    ) -> Optional[LimitUpDiagnosis]:
        """诊断单只涨停股 — 加载数据并做缠论归因"""

        df = self._load_stock_data(symbol, date)
        if df is None or len(df) < 80:
            return None

        n = len(df)
        close = df['close'].values.astype(np.float64)
        high = df['high'].values.astype(np.float64)
        low = df['low'].values.astype(np.float64)
        volume = df['volume'].values.astype(np.float64) if 'volume' in df.columns else np.ones(n)

        # 检查是否涨停 (涨跌幅 >= 9.5%)
        if n < 2:
            return None
        chg = (close[-1] / close[-2] - 1) * 100
        if chg < self.LIMIT_UP_THRESHOLD:
            return None

        # 流动性过滤
        if 'amount' in df.columns:
            avg_to = df['amount'].tail(20).mean()
        elif 'volume' in df.columns:
            avg_to = (df['volume'].tail(20) * df['close'].tail(20)).mean()
        else:
            avg_to = float('inf')
        if avg_to < self.min_volume:
            return None

        # === 计算缠论结构 ===
        ema12 = _calc_ema(close, 12)
        ema26 = _calc_ema(close, 26)
        macd_line = ema12 - ema26
        macd_hist = 2 * (macd_line - _calc_ema(macd_line, 9))
        ema20 = _calc_ema(close, 20)
        ema60 = _calc_ema(close, 60)
        ema120 = _calc_ema(close, 120)

        chan = compute_enhanced_chan_output(
            close, high, low, ema20, ema60, ema120, macd_hist, volume,
        )

        d = LimitUpDiagnosis(symbol=symbol)
        d.name = self._get_stock_name(symbol)
        try:
            raw_ind = self._industry_map.get(d.symbol, '')
            d.industry = get_industry_category(raw_ind) if raw_ind else '未归类'
        except Exception:
            raw_ind = self._industry_map.get(d.symbol, '')
            d.industry = get_industry_category(raw_ind) if raw_ind else '未归类'
        d.close = float(close[-1])
        d.change_pct = float(chg)
        d.volume_ratio = float(volume[-1] / np.mean(volume[-6:-1])) if np.mean(volume[-6:-1]) > 0 else 1.0

        # === 涨停前结构 (5日前) ===
        pre_idx = max(0, n - 6)  # 涨停前1天 (涨停日是n-1)
        d.pre_trend_type = int(chan['trend_type'][pre_idx])
        d.pre_pivot_position = int(chan['pivot_position'][pre_idx])
        d.pre_stroke_direction = int(chan['stroke_direction'][pre_idx])
        d.pre_buy_point = int(chan['buy_point'][pre_idx])
        d.pre_buy_confidence = float(chan['buy_confidence'][pre_idx])
        d.pre_bottom_divergence = bool(chan.get('hidden_bottom_divergence', np.zeros(n))[pre_idx])
        d.pre_alignment_score = float(chan['alignment_score'][pre_idx])

        # 涨停前盘整天数（趋势类型=1的连续K线数）
        cons_days = 0
        for i in range(pre_idx, max(0, pre_idx - 30), -1):
            if chan['trend_type'][i] == 1:
                cons_days += 1
            else:
                break
        d.pre_consolidation_days = cons_days

        # 涨停日信号
        d.cur_buy_point = int(chan['buy_point'][-1])
        d.cur_buy_confidence = float(chan['buy_confidence'][-1])
        d.cur_bottom_divergence = bool(chan.get('hidden_bottom_divergence', np.zeros(n))[-1])

        # 中枢突破检测
        zg = float(chan['pivot_zg'][-1])
        zd = float(chan['pivot_zd'][-1])
        if not np.isnan(zg):
            d.pivot_zg = zg
            d.pivot_zd = zd
            d.above_pivot_pct = (close[-1] - zg) / zg * 100 if zg > 0 else 0
            # 突破判断: 涨停日收盘 > ZG 且 前日收盘 <= ZG
            d.pivot_breakout = (close[-1] > zg and close[-2] <= zg * 1.01)
            if not d.pivot_breakout and close[-1] > zg * 1.03:
                d.pivot_breakout = True  # 大幅突破也计入

        if not np.isnan(zd):
            d.below_pivot_pct = (zd - close[-2]) / zd * 100 if zd > 0 else 0

        # 量能
        avg_vol_5 = np.mean(volume[-6:-1]) if len(volume) >= 6 and np.mean(volume[-6:-1]) > 0 else 1
        d.vol_spike = volume[-1] / avg_vol_5

        # 底分型质量
        d.bottom_fx_quality = float(chan.get('bottom_fractal_quality', np.zeros(n))[-1])

        # === 缠论归因分类 ===
        d.limit_up_type, d.type_confidence, d.chan_explanation = self._classify_limit_up_type(d)

        return d

    def _classify_limit_up_type(self, d: LimitUpDiagnosis) -> Tuple[str, float, str]:
        """基于缠论结构对涨停进行归因分类

        六大涨停模式：
        1. 中枢突破板 — 突破中枢上沿ZG，三买B3触发
        2. 底背驰反转板 — 下跌趋势+底背驰+一买B1 → V型反转
        3. 二买确认板 — 一买后回调不破前低，B2确认后加速涨停
        4. 趋势加速板 — 上涨趋势+中枢上方+放量，延续加速
        5. 笔耗尽反转板 — 下跌笔趋势耗尽+底分型 → 反转涨停
        6. 盘整突破板 — 长期盘整后首日突破+放量
        """
        scores = []  # [(type_name, confidence, explanation)]

        # 1. 中枢突破板
        if d.pivot_breakout and d.above_pivot_pct > 0:
            conf = min(0.95, 0.55 + d.above_pivot_pct / 20 + (0.15 if d.vol_spike > 1.5 else 0))
            scores.append(('中枢突破板', conf,
                f'突破中枢上沿ZG({d.pivot_zg:.2f})，收盘高于ZG {d.above_pivot_pct:.1f}%'
                f'{", 放量" if d.vol_spike > 1.5 else ""}'
                f'{", B3三买触发" if d.cur_buy_point == 3 else ""}'))

        # 2. 底背驰反转板
        if d.pre_bottom_divergence and d.pre_trend_type == -2 and d.pre_pivot_position <= 0:
            conf = 0.55 + (0.15 if d.pre_buy_point == 1 else 0) + (0.1 if d.bottom_fx_quality > 0.3 else 0)
            scores.append(('底背驰反转板', min(0.9, conf),
                f'下跌趋势+底背驰，前5日底背驰确认'
                f'{", B1一买触发" if d.pre_buy_point == 1 else ""}'
                f'{", 底分型质量{:.0%}" if d.bottom_fx_quality > 0.3 else ""}'.format(d.bottom_fx_quality)))

        # 3. 二买确认板
        if d.pre_buy_point == 2 and d.pre_trend_type >= 1 and d.pre_pivot_position <= 0:
            conf = 0.60 + (0.15 if d.vol_spike > 1.5 else 0) + (0.1 if d.pre_alignment_score > 0 else 0)
            scores.append(('二买确认板', min(0.9, conf),
                f'B2二买后加速，前日笔方向{d.pre_stroke_direction}'
                f'{", 均线多头排列" if d.pre_alignment_score > 0.3 else ""}'))

        # 4. 趋势加速板
        if d.pre_trend_type == 2 and d.pre_pivot_position == 1 and d.pre_alignment_score > 0.3:
            conf = 0.50 + (0.15 if d.vol_spike > 2 else 0) + (0.1 if d.cur_buy_point >= 3 else 0)
            scores.append(('趋势加速板', min(0.85, conf),
                f'上涨趋势+中枢上方，均线多头排列'
                f'{", 放量加速" if d.vol_spike > 2 else ""}'))

        # 5. 笔耗尽反转板
        if d.pre_stroke_direction == -1 and d.bottom_fx_quality > 0.25:
            # 前日下跌笔 + 底分型质量 → 笔趋势耗尽
            conf = 0.45 + d.bottom_fx_quality * 0.3 + (0.1 if d.vol_spike > 1.5 else 0)
            scores.append(('笔耗尽反转板', min(0.8, conf),
                f'下跌笔趋势耗尽+底分型(质量{d.bottom_fx_quality:.0%})→反转'
                f'{", 量在价先" if d.vol_spike > 2 else ""}'))

        # 6. 盘整突破板
        if d.pre_consolidation_days >= 8 and d.vol_spike > 1.5:
            conf = 0.45 + min(d.pre_consolidation_days / 50, 0.25) + (0.15 if d.vol_spike > 2 else 0)
            scores.append(('盘整突破板', min(0.85, conf),
                f'盘整{d.pre_consolidation_days}天后突破'
                f'{", 倍量启动" if d.vol_spike > 2 else ", 放量启动"}'))

        if scores:
            scores.sort(key=lambda x: x[1], reverse=True)
            return scores[0]

        # 兜底: 根据现有数据定性
        if d.vol_spike > 2:
            return ('放量异动板', 0.35, '无明显缠论买点，但放量明显，关注后续结构演变')
        return ('未分类', 0.2, f'涨{d.change_pct:.1f}%，量比{d.volume_ratio:.1f}，建议观察后续走势结构')

    def _extract_limit_up_commonalities(
        self, diagnoses: List[LimitUpDiagnosis]
    ) -> List[str]:
        """从涨停股中提取共性特征"""
        traits = []
        n = len(diagnoses)

        # 1. 主导模式统计
        type_counts = defaultdict(int)
        for d in diagnoses:
            type_counts[d.limit_up_type] += 1
        top_type, top_count = max(type_counts.items(), key=lambda x: x[1])
        if top_count >= n * 0.3:
            traits.append(f'今日涨停主导模式: {top_type} ({top_count}/{n}, {top_count/n:.0%})，'
                         f'暗示当前市场偏好此类结构')

        # 2. 涨停前结构共性
        pre_up_trend = sum(1 for d in diagnoses if d.pre_trend_type == 2)
        pre_down_trend = sum(1 for d in diagnoses if d.pre_trend_type == -2)
        pre_cons = sum(1 for d in diagnoses if d.pre_trend_type == 1)

        if pre_up_trend >= n * 0.5:
            traits.append(f'涨停前已处上涨趋势: {pre_up_trend}/{n} ({pre_up_trend/n:.0%})，'
                         f'趋势延续型涨停为主，选股应偏向已形成上涨结构的个股')
        elif pre_down_trend >= n * 0.5:
            traits.append(f'涨停前处下跌趋势: {pre_down_trend}/{n} ({pre_down_trend/n:.0%})，'
                         f'超跌反弹型涨停为主，需警惕次日分化')
        elif pre_cons >= n * 0.5:
            traits.append(f'涨停前处盘整: {pre_cons}/{n} ({pre_cons/n:.0%})，'
                         f'盘整突破型，关注突破后的持续性')

        # 3. 中枢位置共性
        pivot_below = sum(1 for d in diagnoses if d.pre_pivot_position == -1)
        pivot_above = sum(1 for d in diagnoses if d.pre_pivot_position == 1)
        if pivot_below >= n * 0.4:
            traits.append(f'{pivot_below}/{n}涨停股前日处中枢下方，中枢突破是主要驱动力')
        if pivot_above >= n * 0.5:
            traits.append(f'{pivot_above}/{n}涨停股前日处中枢上方，强势股延续，追涨风险相对可控')

        # 4. 量能共性
        high_vol = sum(1 for d in diagnoses if d.vol_spike > 2)
        if high_vol >= n * 0.6:
            traits.append(f'{high_vol}/{n}涨停股倍量启动(量比>2x)，量在价先信号明确，'
                         f'但需关注次日量能是否持续')

        # 5. 均线排列共性
        bullish_align = sum(1 for d in diagnoses if d.pre_alignment_score > 0.3)
        if bullish_align >= n * 0.5:
            traits.append(f'{bullish_align}/{n}涨停前均线多头排列，技术共振增强涨停可靠性')

        # 6. 底背驰共性
        bottom_div = sum(1 for d in diagnoses if d.pre_bottom_divergence or d.cur_bottom_divergence)
        if bottom_div >= n * 0.3:
            traits.append(f'{bottom_div}/{n}涨停股呈现底背驰特征，'
                         f'缠论"没有趋势，没有背驰"→背驰后的反转力度较强')

        # 7. 行业集中
        ind_counts = defaultdict(int)
        for d in diagnoses:
            ind_counts[d.industry] += 1
        top_ind = max(ind_counts, key=ind_counts.get)
        top_ind_count = ind_counts[top_ind]
        if top_ind_count >= n * 0.25:
            traits.append(f'涨停集中在[{top_ind}]({top_ind_count}/{n})，'
                         f'板块效应显著，该行业可能处于主升浪')

        # 8. 综合建议
        if n >= 5:
            traits.append(f'操作建议: 今日{n}只涨停股中，'
                         f'建议关注{top_type}标的（缠论结构最清晰）')

        return traits

    # ========== 涨停板报告打印 ==========

    def _print_limit_up_analysis(self, summary: LimitUpSummary, n: int):
        """打印涨停板缠论归因分析"""
        if not summary.limit_up_stocks:
            return

        print(f'\n┌{"─"*76}┐')
        print(f'│  Step 3: 涨停板缠论归因分析{" " * 50}│')
        print(f'├{"─"*76}┤')

        # 概览
        print(f'│  今日涨停: {summary.total_limit_up} 只{" " * 60}│')
        if summary.dominant_pattern:
            print(f'│  主导模式: {summary.dominant_pattern}{" " * (64 - len(summary.dominant_pattern))}│')

        # 类型分布
        type_items = sorted(summary.type_distribution.items(), key=lambda x: x[1], reverse=True)
        dist_str = '  '.join(f'{t}:{c}只' for t, c in type_items[:6])
        print(f'│  类型分布: {dist_str:<64s}│')

        # 行业集中
        top_inds = sorted(summary.industry_concentration.items(), key=lambda x: x[1], reverse=True)[:5]
        if top_inds:
            ind_str = '  '.join(f'{i}:{c}只' for i, c in top_inds)
            print(f'│  涨停行业集中: {ind_str:<58s}│')

        print(f'├{"─"*76}┤')

        # 个股归因明细
        print(f'│  {"代码":<10s} {"名称":<10s} {"涨跌":>6s} {"量比":>5s} '
              f'{"归因类型":<14s} {"置信":>5s} {"涨停前缠论结构":^30s}│')
        print(f'│  {"-"*74}│')

        trend_map = {2: '上涨趋势', 1: '盘整', 0: '--', -2: '下跌趋势'}
        pp_map = {1: '中枢上', 0: '中枢内', -1: '中枢下'}

        for d in summary.limit_up_stocks[:n]:
            pre_structure = (
                f'{trend_map.get(d.pre_trend_type, "?"):>4s} | '
                f'{pp_map.get(d.pre_pivot_position, "?"):>3s} | '
                f'笔:{"↑" if d.pre_stroke_direction == 1 else "↓" if d.pre_stroke_direction == -1 else "—"}'
            )
            print(f'│  {d.symbol:<10s} {d.name:<10s} {d.change_pct:>+5.1f}% '
                  f'{d.vol_spike:>4.1f}x {d.limit_up_type:<14s} {d.type_confidence:>4.0%} '
                  f'{pre_structure:<30s}│')

        print(f'├{"─"*76}┤')

        # 缠论解释
        print(f'│  缠论归因详情:{" " * 62}│')
        for d in summary.limit_up_stocks[:min(n, 8)]:
            # 截短解释以避免溢出
            expl = d.chan_explanation[:72]
            print(f'│  [{d.symbol} {d.name}] {d.limit_up_type} ({d.type_confidence:.0%})')
            print(f'│    → {expl:<70s}│')

        print(f'└{"─"*76}┘')

        # 共性总结
        if summary.common_traits:
            print(f'\n┌{"─"*76}┐')
            print(f'│  涨停板共性总结{" " * 62}│')
            print(f'├{"─"*76}┤')
            for i, trait in enumerate(summary.common_traits):
                print(f'│  {i+1}. {trait:<72s}│')
            print(f'└{"─"*76}┘')

    # ========== 持仓与选股加载 ==========

    def _load_portfolio_state(self) -> Dict:
        """加载当前持仓状态"""
        pf_file = os.path.join(PROJECT_ROOT, 'trade', 'portfolio_state.json')
        if os.path.exists(pf_file):
            try:
                with open(pf_file) as f:
                    return json.load(f)
            except Exception:
                pass
        return {"cash": 150000.0, "positions": {}}

    def _load_today_selections(self, date_str: str) -> Optional[pd.DataFrame]:
        """加载当日的选股结果（来自export_selection.py输出）"""
        csv_path = os.path.join(PROJECT_ROOT, 'output', '选股历史.csv')
        if not os.path.exists(csv_path):
            return None
        try:
            df = pd.read_csv(csv_path, dtype={'股票代码': str})
            # 筛选当天的记录
            today_sel = df[df['选股日期'] == date_str]
            if len(today_sel) == 0:
                return None
            return today_sel
        except Exception:
            return None

    def _diagnose_portfolio_holdings(
        self, holdings: Dict, diagnoses: List[StockDiagnosis], date: str
    ) -> List[Dict]:
        """评估当前持仓的缠论状态（增强版：含深度技术分析）"""
        diag_map = {d.symbol: d for d in diagnoses}
        result = []
        for code, pos_info in holdings.items():
            if isinstance(pos_info, dict):
                shares = pos_info.get('shares', 0)
                cost = pos_info.get('cost_price', 0)
            else:
                shares = float(pos_info) if pos_info else 0
                cost = 0

            d = diag_map.get(code)
            if d is None:
                result.append({
                    'code': code, 'name': '', 'shares': shares, 'cost': cost,
                    'close': 0, 'change_pct': 0, 'market_value': 0, 'pnl_pct': 0,
                    'trend': '-', 'buy_pt': 0, 'sell_pt': 0, 'composite': 0,
                    'risk': '未知', 'rank': '-', 'stop_loss': 0,
                    'action': '关注', 'reason': '未在诊断范围',
                    'structure_desc': '', 'key_levels': '', 'div_desc': '无数据',
                    'vol_desc': '无数据', 'risk_detail': '', 'take_profit': '',
                    'analysis': '',
                })
                continue

            market_val = shares * d.close if shares and d.close else 0
            pnl_pct = (d.close / cost - 1) * 100 if cost > 0 and d.close > 0 else 0

            # ---- 缠论结构描述 ----
            trend_label = {2: '上涨趋势', 1: '盘整', 0: '-', -2: '下跌趋势'}.get(d.trend_type, '?')
            pivot_label = {-1: '中枢下方', 0: '中枢内部', 1: '中枢上方'}.get(d.pivot_position, '?')
            stroke_dir = '↑' if d.stroke_direction > 0 else '↓' if d.stroke_direction < 0 else '→'
            segment_dir = '↑' if d.segment_direction > 0 else '↓' if d.segment_direction < 0 else '→'

            structure_parts = [f'{trend_label}', f'{pivot_label}',
                              f'笔{stroke_dir}/线段{segment_dir}']
            if d.pivot_zg > 0 and d.pivot_zd > 0:
                structure_parts.append(f'中枢[{d.pivot_zd:.2f}-{d.pivot_zg:.2f}]')
            structure_desc = '，'.join(structure_parts)

            # ---- 买卖点分析 ----
            bp_label = {1: 'B1', 2: 'B2', 3: 'B3'}.get(d.buy_point, '-')
            if d.second_buy:
                bp_label = 'B2*'
            sp_label = f'S{d.sell_point}' if d.sell_point > 0 else '-'
            bp_note = '已确认' if d.confirmed_buy else '潜在' if d.buy_point > 0 else '无'
            sp_note = '已确认' if d.confirmed_sell else '潜在' if d.sell_point > 0 else '无'

            # ---- 背驰分析 ----
            div_parts = []
            if d.bottom_divergence:
                div_parts.append('底背驰✓(下跌力竭)')
            if d.top_divergence:
                div_parts.append('顶背驰⚠(上涨力竭)')
            div_desc = '；'.join(div_parts) if div_parts else '无背驰信号'

            # ---- 量价分析 ----
            vol_desc = d.volume_anomaly if d.volume_anomaly else ('放量' if d.volume_ratio > 1.5 else '缩量' if d.volume_ratio < 0.7 else '正常')

            # ---- 关键价位 ----
            levels = []
            if d.stop_loss_price > 0:
                levels.append(f'止损{d.stop_loss_price:.2f}')
            if d.pivot_zg > 0:
                levels.append(f'中枢上沿{d.pivot_zg:.2f}')
            if d.pivot_zd > 0:
                levels.append(f'中枢下沿{d.pivot_zd:.2f}')
            if cost > 0:
                levels.append(f'成本{cost:.2f}')
            key_levels = ' | '.join(levels)

            # ---- 风险判断 ----
            risk_reasons = []
            if d.ma_breakdowns:
                risk_reasons.append(f'均线破位({",".join(d.ma_breakdowns)})')
            if d.pivot_breakdown:
                risk_reasons.append('中枢下方运行')
            if d.top_divergence:
                risk_reasons.append('顶背驰')
            if d.confirmed_sell:
                risk_reasons.append('卖点确认')
            risk_detail = '；'.join(risk_reasons) if risk_reasons else '无明显风险信号'

            # ---- 止盈参考 ----
            take_profit = ''
            if d.pivot_zg > 0 and d.close > d.pivot_zg:
                take_profit = f'中枢上沿{d.pivot_zg:.2f}已突破，看前高'
            elif d.pivot_zd > 0 and d.close < d.pivot_zd:
                take_profit = f'需先收复中枢下沿{d.pivot_zd:.2f}'

            # ---- 操作建议 ----
            if d.confirmed_sell or d.sell_point >= 1:
                action = '⚠ 卖出'
                reason = f'出现卖点信号(S{d.sell_point})'
            elif d.composite_score < -0.2:
                action = '⚠ 减仓'
                reason = f'综合评分偏低({d.composite_score:+.2f})'
            elif d.confirmed_buy and d.buy_strength > 0.5:
                action = '持盈'
                reason = f'买点确认，强度{d.buy_strength:.2f}'
            elif d.buy_point > 0 and d.composite_score > 0:
                action = '持有'
                reason = f'有买点支撑(B{d.buy_point})'
            elif d.trend_type == 2:
                action = '持有'
                reason = '上涨趋势中'
            else:
                action = '关注'
                reason = '趋势不明朗'

            # ---- 综合分析文本 ----
            analysis_lines = [
                f'{d.name}({code}): 收盘{d.close:.2f}({d.change_pct:+.1f}%), '
                f'持仓{shares}股, 成本{cost:.2f}, '
                f'市值{market_val:,.0f}, 盈亏{pnl_pct:+.1f}%',
                f'缠论结构: {structure_desc}',
                f'买卖点: 买点={bp_label}({bp_note}), 卖点={sp_label}({sp_note})',
                f'背驰: {div_desc}',
                f'量价: 量比{d.volume_ratio:.1f}x({vol_desc})',
                f'综合评分: {d.composite_score:+.2f} ({d.opportunity_rank}级), '
                f'风险: {d.risk_level}',
                f'关键价位: {key_levels}',
                f'操作建议: {action} — {reason}',
            ]
            if risk_detail:
                analysis_lines.append(f'风险提示: {risk_detail}')
            if take_profit:
                analysis_lines.append(f'止盈参考: {take_profit}')

            result.append({
                'code': code, 'name': d.name, 'shares': shares, 'cost': cost,
                'close': d.close, 'change_pct': d.change_pct,
                'market_value': market_val, 'pnl_pct': pnl_pct,
                'trend': trend_label, 'buy_pt': bp_label, 'sell_pt': sp_label,
                'composite': d.composite_score, 'risk': d.risk_level,
                'rank': d.opportunity_rank, 'stop_loss': d.stop_loss_price,
                'action': action, 'reason': reason,
                'structure_desc': structure_desc,
                'key_levels': key_levels,
                'div_desc': div_desc,
                'vol_desc': f'量比{d.volume_ratio:.1f}x({vol_desc})',
                'risk_detail': risk_detail,
                'take_profit': take_profit,
                'analysis': '\n'.join(analysis_lines),
            })
        return result

    def _generate_replacement_plan(
        self, holdings_diag: List[Dict], today_selections: List[Dict]
    ) -> List[Dict]:
        """生成持仓替换计划 — 走坏持仓 → 当日选股评分排名依次替换"""
        danger = [h for h in holdings_diag if h.get('action', '') == '⚠ 卖出']
        warn = [h for h in holdings_diag if h.get('action', '') == '⚠ 减仓']
        plan = []

        if not danger and not warn:
            return plan

        hold_codes = {h['code'] for h in holdings_diag}
        available = [s for s in today_selections
                     if str(s.get('股票代码', '')) not in hold_codes]
        available.sort(key=lambda x: float(x.get('综合评分', 0)), reverse=True)

        for h in danger:
            if not available:
                break
            r = available.pop(0)
            plan.append({
                'code': h['code'], 'name': h['name'],
                'reason': f"结构恶化({h.get('sell_pt', '卖点')}/{h.get('trend', '趋势反转')})",
                'action': '卖出',
                'replacement_code': str(r.get('股票代码', '')),
                'replacement_name': str(r.get('股票名称', '')),
                'replacement_score': float(r.get('综合评分', 0)),
                'replacement_price': float(r.get('当前价格', 0)),
                'replacement_signal': str(r.get('信号强度', '')),
                'urgency': '立即',
            })

        for h in warn:
            if not available:
                break
            r = available.pop(0)
            plan.append({
                'code': h['code'], 'name': h['name'],
                'reason': f"评分偏低/趋势不明({h.get('composite', 0):+.2f})",
                'action': '预备替换',
                'replacement_code': str(r.get('股票代码', '')),
                'replacement_name': str(r.get('股票名称', '')),
                'replacement_score': float(r.get('综合评分', 0)),
                'replacement_price': float(r.get('当前价格', 0)),
                'replacement_signal': str(r.get('信号强度', '')),
                'urgency': '明日确认',
            })

        return plan

    # ========== Step 6: 作战计划 + 系统进化 ==========

    def generate_portfolio_plan(
        self, diagnoses: List[StockDiagnosis],
        market: MarketVitalSigns, today_selections: List[Dict] = None,
        max_positions: int = 8
    ) -> List[StockDiagnosis]:
        """生成作战计划 — 对齐策略选股结果+缠论诊断交叉分析"""
        diag_map = {d.symbol: d for d in diagnoses}
        result = []

        if today_selections:
            # 使用策略选股结果, 交叉缠论诊断做深度分析
            for s in today_selections:
                code = str(s.get('股票代码', ''))
                d = diag_map.get(code)
                if d:
                    # 用策略的评分替代纯缠论评分, 但保留缠论结构数据
                    d.composite_score = float(s.get('综合评分', d.composite_score))
                    d.opportunity_rank = self._score_to_rank(d.composite_score)
                else:
                    # 策略选了但缠论诊断未覆盖(如北交所), 创建基础诊断
                    d = StockDiagnosis(
                        symbol=code,
                        name=str(s.get('股票名称', '')),
                        industry=str(s.get('行业', '')),
                        close=float(s.get('当前价格', 0)),
                        composite_score=float(s.get('综合评分', 0)),
                        opportunity_rank=self._score_to_rank(float(s.get('综合评分', 0))),
                    )
                result.append(d)
        else:
            # 无策略选股数据时回退: 按缠论评分独立筛选
            if market.sentiment_cycle == '冰点':
                max_positions = max(2, max_positions // 4)
            elif market.sentiment_cycle == '高潮':
                max_positions = min(max_positions, 8)
            elif market.sentiment_cycle == '退潮':
                max_positions = max(1, max_positions // 3)
            elif market.bear_risk:
                max_positions = max(1, max_positions // 2)

            candidates = [d for d in diagnoses
                          if d.composite_score >= 0.15 and d.opportunity_rank in ('A', 'B')]
            candidates.sort(key=lambda x: x.composite_score, reverse=True)

            used_industries = set()
            for d in candidates:
                if len(result) >= max_positions:
                    break
                if len(result) < 3 or d.industry not in used_industries:
                    result.append(d)
                    used_industries.add(d.industry)

        return result

    @staticmethod
    def _score_to_rank(score: float) -> str:
        if score >= 0.5: return 'A'
        if score >= 0.2: return 'B'
        if score >= 0.0: return 'C'
        return 'D'

    def _get_previous_trading_date(self, current_date: str) -> Optional[str]:
        """从输出目录找到上一个复盘日期"""
        import glob
        pattern = os.path.join(self.output_dir, '缠论全量分析_*.csv')
        files = glob.glob(pattern)
        dates = set()
        for f in files:
            bn = os.path.basename(f)
            # 缠论全量分析_YYYYMMDD.csv
            d = bn.replace('缠论全量分析_', '').replace('.csv', '')
            if len(d) == 8 and d.isdigit():
                dates.add(d)
        dates = sorted(dates, reverse=True)
        target = current_date.replace('-', '')
        for d in dates:
            if d < target:
                return d
        return None

    def _evaluate_prediction_accuracy(
        self, current_date: str
    ) -> Dict:
        """
        加载上个交易日的缠论全量分析表，对比今日实际涨跌，评估预测准确性。

        核心指标:
        - direction_accuracy: 买入信号的方向正确率（推荐买入后涨了）
        - avg_return: 推荐标的平均涨跌幅
        - by_type: 各买点类型(B1/B2/B3)的准确率
        - by_confidence: 不同置信度区间的准确率
        - market_return: 同期市场平均涨跌幅（基准）
        """
        prev_date_str = self._get_previous_trading_date(current_date)
        result = {
            'available': False,
            'prev_date': prev_date_str,
            'direction_accuracy': None,
            'avg_return': None,
            'market_return': None,
            'by_type': {},
            'by_confidence': {},
            'total_buy_signals': 0,
            'hits': 0,
        }

        if not prev_date_str:
            return result

        prev_csv = os.path.join(self.output_dir, f'缠论全量分析_{prev_date_str}.csv')
        if not os.path.exists(prev_csv):
            return result

        try:
            prev_df = pd.read_csv(prev_csv, dtype={'股票代码': str})
        except Exception:
            return result

        # 筛选有买点信号的股票
        buy_df = prev_df[prev_df['买点类型'].notna() & (prev_df['买点类型'] != '-') & (prev_df['买点类型'] != '')]
        if len(buy_df) == 0:
            return result

        result['total_buy_signals'] = len(buy_df)

        # 为每只股票获取今日涨跌幅
        hits = 0
        returns = []
        by_type_hits = defaultdict(lambda: {'total': 0, 'hits': 0})
        by_conf_hits = defaultdict(lambda: {'total': 0, 'hits': 0})
        market_returns = []

        for _, row in buy_df.iterrows():
            code = row['股票代码']
            prev_close = row.get('收盘价', 0)
            if pd.isna(prev_close) or prev_close <= 0:
                continue

            # 加载今日数据
            sym = code
            if not code.startswith('sh') and not code.startswith('sz'):
                for prefix in ['sh', 'sz']:
                    test_sym = f'{prefix}{code}'
                    if os.path.exists(os.path.join(self.data_dir, f'{test_sym}_qfq.csv')):
                        sym = test_sym
                        break

            df = self._load_stock_data(sym, current_date, lookback=5)
            if df is None or len(df) < 2:
                continue

            today_close = float(df['close'].iloc[-1])
            ret_pct = (today_close / prev_close - 1) * 100
            returns.append(ret_pct)

            # 方向准确率: 买入信号的股票今天涨了
            if ret_pct > 0:
                hits += 1

            # 按买点类型
            buy_type = str(row.get('买点类型', ''))
            if buy_type and buy_type != '-':
                by_type_hits[buy_type]['total'] += 1
                if ret_pct > 0:
                    by_type_hits[buy_type]['hits'] += 1

            # 按置信度区间
            conf = row.get('买点置信度', 0)
            if pd.notna(conf) and conf > 0:
                if conf >= 0.7:
                    bucket = 'high(>=0.7)'
                elif conf >= 0.5:
                    bucket = 'mid(0.5-0.7)'
                else:
                    bucket = 'low(<0.5)'
            else:
                bucket = 'unknown'
            by_conf_hits[bucket]['total'] += 1
            if ret_pct > 0:
                by_conf_hits[bucket]['hits'] += 1

        # 市场基准: 随机抽100只非推荐股票
        all_symbols = self._list_available_stocks()
        np.random.seed(42)
        sample_symbols = np.random.choice(all_symbols, min(100, len(all_symbols)), replace=False)
        for sym in sample_symbols:
            df = self._load_stock_data(sym, current_date, lookback=5)
            if df is not None and len(df) >= 2:
                mr = (float(df['close'].iloc[-1]) / float(df['close'].iloc[-2]) - 1) * 100
                market_returns.append(mr)

        result['available'] = True
        result['hits'] = hits
        total_evaluated = len(returns)
        result['direction_accuracy'] = hits / max(total_evaluated, 1)
        result['avg_return'] = float(np.mean(returns)) if returns else 0.0
        result['market_return'] = float(np.mean(market_returns)) if market_returns else 0.0
        result['excess_return'] = result['avg_return'] - result['market_return']

        for bt, d in by_type_hits.items():
            result['by_type'][bt] = {
                'accuracy': d['hits'] / max(d['total'], 1),
                'count': d['total'],
            }
        for bk, d in by_conf_hits.items():
            result['by_confidence'][bk] = {
                'accuracy': d['hits'] / max(d['total'], 1),
                'count': d['total'],
            }

        return result

    def _update_accuracy_history(self, pred_accuracy: Dict, date: str):
        """追加预测准确性到历史记录文件"""
        if not pred_accuracy or not pred_accuracy.get('available'):
            return
        history_path = os.path.join(self.output_dir, 'accuracy_history.json')
        history = []
        if os.path.exists(history_path):
            try:
                with open(history_path, 'r') as f:
                    history = json.load(f)
            except Exception:
                history = []
        # 去重：同一天不重复记录
        existing_dates = {h.get('date', '') for h in history}
        if date in existing_dates:
            return
        history.append({
            'date': date,
            'prev_date': pred_accuracy.get('prev_date', ''),
            'direction_accuracy': round(pred_accuracy['direction_accuracy'], 4),
            'avg_return': round(pred_accuracy['avg_return'], 4),
            'market_return': round(pred_accuracy.get('market_return', 0), 4),
            'excess_return': round(pred_accuracy.get('excess_return', 0), 4),
            'total_signals': pred_accuracy['total_buy_signals'],
            'hits': pred_accuracy['hits'],
            'by_type': pred_accuracy.get('by_type', {}),
        })
        # 只保留最近60条
        history = history[-60:]
        with open(history_path, 'w') as f:
            json.dump(history, f, ensure_ascii=False, indent=2)

    def _load_accuracy_history(self) -> List[Dict]:
        """加载预测准确性历史"""
        history_path = os.path.join(self.output_dir, 'accuracy_history.json')
        if os.path.exists(history_path):
            try:
                with open(history_path, 'r') as f:
                    return json.load(f)
            except Exception:
                pass
        return []

    def generate_evolution_patches(
        self, diagnoses: List[StockDiagnosis],
        market: MarketVitalSigns, industries: List[IndustryMainLine],
        pred_accuracy: Dict = None,
    ) -> List[EvolutionPatch]:
        """生成系统自进化补丁 — 结合复盘发现 + 预测准确性反馈"""

        patches = []

        # ═══ 数据驱动的进化（基于预测准确性回测） ═══
        if pred_accuracy and pred_accuracy.get('available'):
            acc = pred_accuracy['direction_accuracy']
            avg_ret = pred_accuracy['avg_return']
            excess = pred_accuracy.get('excess_return', 0)
            prev_date = pred_accuracy.get('prev_date', '?')

            # 方向准确率过低 → 收紧买入阈值
            if acc is not None and acc < 0.40:
                patches.append(EvolutionPatch(
                    section='signal', key='buy_threshold',
                    current_value=0.06, suggested_value=round(min(0.10, 0.06 + (0.40 - acc) * 0.15), 2),
                    reason=f'[{prev_date}]预测方向准确率仅{acc:.0%}({pred_accuracy["hits"]}/{pred_accuracy["total_buy_signals"]})，'
                           f'远低于50%随机基准，急需收紧买入阈值以过滤噪音信号',
                    confidence=min(0.95, 0.6 + (0.40 - acc)), urgency='urgent',
                ))
            elif acc is not None and acc < 0.50:
                patches.append(EvolutionPatch(
                    section='signal', key='buy_threshold',
                    current_value=0.06, suggested_value=round(0.06 + (0.50 - acc) * 0.10, 2),
                    reason=f'[{prev_date}]预测方向准确率{acc:.0%}略低于50%，建议小幅提高买入阈值',
                    confidence=0.55, urgency='medium',
                ))
            elif acc is not None and acc > 0.65:
                patches.append(EvolutionPatch(
                    section='signal', key='buy_threshold',
                    current_value=0.06, suggested_value=round(max(0.03, 0.06 - (acc - 0.65) * 0.10), 2),
                    reason=f'[{prev_date}]预测方向准确率{acc:.0%}显著高于随机，当前信号质量好，可适当放宽阈值捕获更多机会',
                    confidence=0.60, urgency='medium',
                ))

            # 超额收益
            if excess < -1.0:
                patches.append(EvolutionPatch(
                    section='portfolio', key='target_volatility',
                    current_value=0.15, suggested_value=0.12,
                    reason=f'[{prev_date}]推荐标的超额收益{excess:+.1f}%，跑输市场，降低目标波动率控制风险',
                    confidence=0.55, urgency='medium',
                ))
            elif excess > 2.0:
                patches.append(EvolutionPatch(
                    section='portfolio', key='max_position',
                    current_value=20, suggested_value=25,
                    reason=f'[{prev_date}]推荐标的超额收益{excess:+.1f}%，显著跑赢市场，可扩大持仓上限',
                    confidence=0.55, urgency='low',
                ))

            # 按买点类型的准确率
            by_type = pred_accuracy.get('by_type', {})
            for bt, info in by_type.items():
                if info['count'] >= 3:
                    if info['accuracy'] < 0.35:
                        # 该类型买点准确率低，降低权重
                        key_map = {'B1(一买)': 'b1_buy_mult', 'B2(二买)': 'b2_buy_mult', 'B3(三买)': 'b3_buy_mult'}
                        mult_key = key_map.get(bt, '')
                        if mult_key:
                            patches.append(EvolutionPatch(
                                section='chan_theory_enhanced.buy_sell_points', key=mult_key,
                                current_value=1.20, suggested_value=0.90,
                                reason=f'[{prev_date}]{bt}准确率仅{info["accuracy"]:.0%}({info["count"]}次)，信号质量差，降低权重',
                                confidence=0.65, urgency='high',
                            ))
                    elif info['accuracy'] > 0.70:
                        key_map = {'B1(一买)': 'b1_buy_mult', 'B2(二买)': 'b2_buy_mult', 'B3(三买)': 'b3_buy_mult'}
                        mult_key = key_map.get(bt, '')
                        if mult_key:
                            patches.append(EvolutionPatch(
                                section='chan_theory_enhanced.buy_sell_points', key=mult_key,
                                current_value=1.20, suggested_value=1.40,
                                reason=f'[{prev_date}]{bt}准确率{info["accuracy"]:.0%}({info["count"]}次)，信号质量优异，提高权重',
                                confidence=0.60, urgency='medium',
                            ))

            # 置信度校准: 高置信是否真的更准？
            by_conf = pred_accuracy.get('by_confidence', {})
            high_acc = by_conf.get('high(>=0.7)', {}).get('accuracy', 0)
            low_acc = by_conf.get('low(<0.5)', {}).get('accuracy', 0)
            if high_acc and low_acc and high_acc < low_acc:
                patches.append(EvolutionPatch(
                    section='chan_theory_enhanced.multi_level', key='min_confidence',
                    current_value=0.40, suggested_value=0.55,
                    reason=f'[{prev_date}]高置信信号({high_acc:.0%})反而比低置信({low_acc:.0%})差，'
                           f'置信度评分体系需重新校准',
                    confidence=0.50, urgency='medium',
                ))

        # ═══ 启发式进化（基于当前市场结构） ═══

        # 买卖信号统计
        buy_count = sum(1 for d in diagnoses if d.buy_point > 0)

        # 1. 买卖信号浓度 → 调整 buy_threshold (仅当无准确率数据时，避免重复)
        if not (pred_accuracy and pred_accuracy.get('available')):
            buy_rate = buy_count / max(len(diagnoses), 1)
            if buy_rate < 0.05:
                patches.append(EvolutionPatch(
                    section='signal', key='buy_threshold',
                    current_value=0.06, suggested_value=0.04,
                    reason=f'买点信号稀少({buy_rate:.1%})，建议降低买入阈值以捕获更多机会',
                    confidence=0.7, urgency='high',
                ))
            elif buy_rate > 0.30:
                patches.append(EvolutionPatch(
                    section='signal', key='buy_threshold',
                    current_value=0.06, suggested_value=0.08,
                    reason=f'买点信号过多({buy_rate:.1%})，建议提高阈值以提升信号质量',
                    confidence=0.65, urgency='medium',
                ))

        # 2. 信号类型分布 → 调整缠论买卖点乘数
        b1_count = sum(1 for d in diagnoses if d.buy_point == 1)
        b2_count = sum(1 for d in diagnoses if d.second_buy)
        b3_count = sum(1 for d in diagnoses if d.buy_point == 3)
        s1_count = sum(1 for d in diagnoses if d.sell_point == 1)

        if b2_count > b1_count:
            patches.append(EvolutionPatch(
                section='chan_theory_enhanced.buy_sell_points', key='b2_buy_mult',
                current_value=1.20, suggested_value=1.30,
                reason=f'二买信号({b2_count})比一买({b1_count})更活跃，提高二买权重',
                confidence=0.6, urgency='medium',
            ))

        if s1_count > buy_count * 0.5:
            patches.append(EvolutionPatch(
                section='chan_theory_enhanced.buy_sell_points', key='s1_sell_mult',
                current_value=0.60, suggested_value=0.50,
                reason=f'一卖信号({s1_count})较多，加强卖出削弱力度',
                confidence=0.55, urgency='medium',
            ))

        # 3. 中阴状态过多 → 调整中阴惩罚
        zhongyin_count = sum(1 for d in diagnoses if hasattr(d, 'zhongyin') and getattr(d, 'zhongyin', False))
        if zhongyin_count > len(diagnoses) * 0.3:
            patches.append(EvolutionPatch(
                section='chan_theory_enhanced.signal', key='zhongyin_penalty',
                current_value=0.85, suggested_value=0.75,
                reason=f'中阴股票占比高({zhongyin_count/len(diagnoses):.0%})，加强中阴惩罚',
                confidence=0.6, urgency='low',
            ))

        # 4. 市场状态 → 调整 regime_multiplier
        if market.bear_risk and market.regime == '熊市':
            patches.append(EvolutionPatch(
                section='regime_multiplier', key='bear',
                current_value=0.80, suggested_value=0.70,
                reason=f'熊市风险警报，建议进一步降低仓位',
                confidence=0.7, urgency='high',
            ))

        # 5. 行业主线检测 → 调整行业因子
        top_industries = [ind for ind in industries if ind.hot_rank <= 3 and ind.total_buy_signals >= 3]
        if top_industries:
            for ind in top_industries:
                patches.append(EvolutionPatch(
                    section=f'industry_factors.{ind.industry}', key='enabled',
                    current_value=1, suggested_value=1,
                    reason=f'强主线行业({ind.hot_rank}): {ind.summary}，建议确保该行业因子配置充足',
                    confidence=0.5, urgency='low',
                ))

        return patches

    # ========== 主运行流程 ==========

    def _auto_update_data_and_select(self, date: str):
        """自动更新股票数据 + 运行选股策略"""
        import subprocess
        print('=' * 50)
        print('[自动] Step 0a: 增量更新股票数据...')
        print('=' * 50)
        try:
            subprocess.run([sys.executable, 'data/data_manager.py'],
                          cwd=PROJECT_ROOT, check=False)
        except Exception as e:
            print(f'  数据更新出错(继续): {e}')

        print()
        print('=' * 50)
        print('[自动] Step 0b: 运行选股策略...')
        print('=' * 50)
        try:
            subprocess.run(
                [sys.executable, 'export_selection.py', '--skip-data-refresh'],
                cwd=PROJECT_ROOT, check=False)
        except Exception as e:
            print(f'  选股运行出错(继续): {e}')
        print()
        print('数据与选股准备完成\n')

    def run(self, date: str = None, top_n: int = 15,
            symbols: List[str] = None, evolve: bool = False,
            llm_report: bool = False, auto_evolve: bool = False,
            auto_update: bool = False) -> DailyReviewReport:
        """运行完整复盘 + 可选进化 + 可选LLM报告

        auto_update=True: 自动更新数据+运行选股 → 一条命令完成全流程
        """
        if date is None:
            date = self._find_latest_date()

        self._llm_enabled = llm_report

        # ── 自动数据更新 + 选股 ──
        if auto_update:
            self._auto_update_data_and_select(date)

        # 获取股票列表
        if symbols is None:
            symbols = self._list_available_stocks()
        print(f'[加载] {len(symbols)} 只股票')

        index_df = self._load_index_data(date)

        # Step 1: 大盘体检
        print('[Step 1/6] 大盘环境深度体检...')
        market = self.examine_market_vitals(date, symbols)

        # Step 2+3+5: 个股全面诊断（缠论+排雷）
        print('[Step 2-5/6] 个股缠论诊断 & 排雷...')
        diagnoses = []
        errors = 0
        for i, sym in enumerate(symbols):
            if (i + 1) % 300 == 0:
                print(f'  进度: {i+1}/{len(symbols)}')
            try:
                d = self.diagnose_stock(sym, date, index_df)
                if d is not None:
                    diagnoses.append(d)
            except Exception as e:
                errors += 1
                if errors <= 3:
                    print(f'  [WARN] {sym}: {e}')

        print(f'  有效诊断: {len(diagnoses)} 只 (错误: {errors})')

        # 行业主线
        print('[Step 2/6] 行业主线分析...')
        industries = self.analyze_industry_main_lines(diagnoses)

        # 提前加载选股结果 (后续机会分类 + 持仓替换 + 作战计划共用)
        print('[持有] 加载持仓与选股结果...')
        pf_state = self._load_portfolio_state()
        portfolio_cash = pf_state.get('cash', 150000.0)
        holdings = pf_state.get('positions', {})
        holdings_diag = self._diagnose_portfolio_holdings(holdings, diagnoses, date)
        today_sel_df = self._load_today_selections(date)
        today_selections = []
        if today_sel_df is not None and len(today_sel_df) > 0:
            today_selections = today_sel_df.to_dict('records')
        print(f'  加载当日选股: {len(today_selections)} 只')

        # 机会 & 风险 (优先使用策略选股结果)
        opportunities = self.classify_opportunities(diagnoses, today_selections)
        risks = self.classify_risks(diagnoses)

        # Step 3: 涨停板缠论归因分析
        print('[Step 3/6] 涨停板缠论归因分析...')
        limit_up_summary = self.analyze_limit_up_stocks(date, symbols, index_df)
        print(f'  涨停股: {limit_up_summary.total_limit_up} 只'
              + (f', 主导模式: {limit_up_summary.dominant_pattern}' if limit_up_summary.dominant_pattern else ''))

        # 生成替换建议
        replacement_plan = self._generate_replacement_plan(holdings_diag, today_selections)

        # Step 6: 作战计划 + 进化
        print('[Step 6/6] 生成作战计划 + 系统进化建议...')
        portfolio = self.generate_portfolio_plan(diagnoses, market, today_selections)

        # 评估上一个交易日的预测准确性（数据驱动进化）
        # --evolve时才运行评估（耗时操作），但历史记录始终加载显示
        pred_accuracy = self._evaluate_prediction_accuracy(date) if evolve else {}
        if pred_accuracy.get('available'):
            print(f'  预测准确性: {pred_accuracy["direction_accuracy"]:.0%} '
                  f'({pred_accuracy["hits"]}/{pred_accuracy["total_buy_signals"]}), '
                  f'平均收益{pred_accuracy["avg_return"]:+.1f}%')
            # 持久化记录准确率历史
            self._update_accuracy_history(pred_accuracy, date)

        accuracy_history = self._load_accuracy_history()
        if accuracy_history:
            print(f'  准确率历史: {len(accuracy_history)} 条记录')

        # 始终运行启发式进化（廉价），--evolve时叠加准确率驱动的进化
        patches = self.generate_evolution_patches(
            diagnoses, market, industries, pred_accuracy=pred_accuracy,
        )

        # 组装报告
        report = DailyReviewReport(
            date=date,
            market=market,
            main_lines=industries,
            opportunities=opportunities,
            risk_alerts=risks,
            portfolio_suggestions=portfolio,
            limit_up_summary=limit_up_summary,
            evolution_patches=patches,
            all_diagnoses=diagnoses,
            pred_accuracy=pred_accuracy,
            portfolio_holdings=holdings_diag,
            today_selections=today_selections,
            replacement_plan=replacement_plan,
            portfolio_cash=portfolio_cash,
            accuracy_history=accuracy_history,
            summary={
                'total_stocks': len(diagnoses),
                'buy_signals': sum(1 for d in diagnoses if d.buy_point > 0),
                'sell_signals': sum(1 for d in diagnoses if d.sell_point > 0),
                'a_rank': sum(1 for d in diagnoses if d.opportunity_rank == 'A'),
                'b_rank': sum(1 for d in diagnoses if d.opportunity_rank == 'B'),
                'up_trend': sum(1 for d in diagnoses if d.trend_type == 2),
                'down_trend': sum(1 for d in diagnoses if d.trend_type == -2),
                'consolidation': sum(1 for d in diagnoses if d.trend_type == 1),
                'confirmed_buy': sum(1 for d in diagnoses if d.confirmed_buy),
                'confirmed_sell': sum(1 for d in diagnoses if d.confirmed_sell),
                'second_buy': sum(1 for d in diagnoses if d.second_buy),
                'high_risk': sum(1 for d in diagnoses if d.risk_level == '高危'),
            },
        )

        # 打印报告
        self._print_report(report, top_n)

        # 导出
        self._export_report(report)

        # K线图不再自动生成（mplfinance图表不准确，用户自行截图）
        chart_files = []

        # LLM叙事报告（API Key可用时自动生成）
        if llm_report or self._get_llm_api_key():
            llm_path = self._export_llm_report(report, date_str=report.date.replace('-', ''),
                                               chart_files=chart_files)
            if llm_path:
                print(f'  公众号版报告: {llm_path}')

        # 进化
        if evolve and patches:
            self._apply_evolution(patches, date, auto_confirm=auto_evolve)

        return report

    # ========== 报告打印 ==========

    def _print_report(self, report: DailyReviewReport, n: int):
        """打印完整复盘报告"""
        m = report.market
        summary = report.summary

        # ═══════════ 头部 ═══════════
        print(f'\n{"="*78}')
        print(f'  量化系统每日复盘报告')
        print(f'  日期: {report.date}')
        print(f'{"="*78}')

        # ─── Step 1: 大盘体检 ───
        print(f'\n┌{"─"*76}┐')
        print(f'│  Step 1: 大盘环境深度体检{" " * 49}│')
        print(f'├{"─"*76}┤')
        print(f'│  上证指数: {m.index_close:>10.1f}  ({m.index_change_pct:+.2f}%)  '
              f'{m.index_volume:<4s}  ({m.volume_ratio:.1f}x均量){" " * 5}│')
        print(f'│  市场状态: {m.regime:<4s} (置信度 {m.regime_confidence:.0%})  |  '
              f'动量: {m.momentum_score:+.2f}  |  波动率: {m.volatility_pct:.1%}  │')
        print(f'│  情绪周期: {m.sentiment_cycle:<4s}  |  '
              f'风格: {m.style_regime:<12s}  |  '
              f'{"⚠ 熊市风险" if m.bear_risk else "✓ 无极端风险"}         │')
        print(f'│  赚钱效应: 涨 {m.advancing} / 跌 {m.declining} '
              f'(上涨率 {m.up_pct:.0f}%)  |  '
              f'涨停 {m.limit_up_count} / 跌停 {m.limit_down_count}'
              f'{" (涨停主导)" if m.limit_up_ratio >= 3 else " (跌停潮)" if m.limit_up_ratio <= 0.3 else ""}│')
        if m.index_divergence:
            print(f'│  ⚠ 指数分化: 上证/深成/创业板 走势不同步                     │')
        print(f'│  综合判定: {m.summary:<62s}│')
        print(f'└{"─"*76}┘')

        # ─── Step 2: 行业主线 ───
        print(f'\n┌{"─"*76}┐')
        print(f'│  Step 2: 板块与主线逻辑拆解{" " * 51}│')
        print(f'├{"─"*76}┤')
        print(f'│  {"行业":<16s} {"股票数":>5s} {"涨趋势":>6s} {"跌趋势":>6s} '
              f'{"B1/B2/B3":>10s} {"梯队":>5s} {"综合":>6s} {"评级":^14s}│')
        print(f'│  {"-"*74}│')

        for ind in report.main_lines[:15]:
            formation_bar = '★' * int(ind.formation_score * 3) + '☆' * (3 - int(ind.formation_score * 3))
            print(f'│  {ind.industry:<16s} {ind.stock_count:>5d} '
                  f'{ind.up_trend_pct:>5.0f}% {ind.down_trend_pct:>5.0f}% '
                  f'{ind.b1_count}/{ind.b2_count}/{ind.b3_count}{"":>4s} '
                  f'{formation_bar:<5s} {ind.avg_composite_score:>+5.2f} '
                  f'{ind.summary[:14]:<14s}│')
        print(f'└{"─"*76}┘')

        # ─── Step 3: 涨停板缠论归因 ───
        if report.limit_up_summary:
            self._print_limit_up_analysis(report.limit_up_summary, n)

        # ─── Step 3+5: 缠论机会 ───
        self._print_signal_section(report.opportunities, 'Step 3+5: 缠论结构扫描 — 买点机会', n, is_buy=True)

        # ─── Step 4: 风险预警 ───
        self._print_signal_section(report.risk_alerts, 'Step 4: 风险预警 — 持仓排雷', n, is_buy=False)

        # ─── Step 6: 作战计划 ───
        print(f'\n┌{"─"*76}┐')
        print(f'│  Step 6: 明日作战计划{" " * 56}│')
        print(f'├{"─"*76}┤')

        if report.portfolio_suggestions:
            # 仓位建议
            if report.market.sentiment_cycle in ('冰点', '退潮'):
                pos_advice = '防御仓位 (1-3成)'
            elif report.market.sentiment_cycle == '高潮':
                pos_advice = '积极仓位 (5-7成)'
            elif report.market.bear_risk:
                pos_advice = '轻仓或空仓 (0-2成)'
            else:
                pos_advice = '中性仓位 (3-5成)'

            print(f'│  仓位建议: {pos_advice:<60s}│')
            print(f'│  推荐标的:{" " * 66}│')
            print(f'│  {"#":>2s} {"代码":<10s} {"名称":<10s} {"行业":<14s} '
                  f'{"收盘":>7s} {"买点":>5s} {"止损":>7s} {"综合":>6s} {"等级":>4s}│')
            for i, d in enumerate(report.portfolio_suggestions):
                bp_label = {1: 'B1', 2: 'B2', 3: 'B3'}.get(d.buy_point, '')
                if d.second_buy:
                    bp_label = 'B2*'
                print(f'│  {i+1:>2d} {d.symbol:<10s} {d.name:<10s} {d.industry:<14s} '
                      f'{d.close:>7.2f} {bp_label:>5s} {d.stop_loss_price:>7.2f} '
                      f'{d.composite_score:>+5.2f} {d.opportunity_rank:>4s}│')
        else:
            print(f'│  无符合条件的推荐标的{" " * 54}│')
        print(f'└{"─"*76}┘')

        # ─── 进化补丁 ───
        if report.evolution_patches:
            print(f'\n┌{"─"*76}┐')
            print(f'│  🔧 系统自进化建议{" " * 60}│')
            print(f'├{"─"*76}┤')
            for p in report.evolution_patches:
                urgency_mark = {'urgent': '🔴', 'high': '🟠', 'medium': '🟡', 'low': '🟢'}
                print(f'│  {urgency_mark.get(p.urgency, "⚪")} [{p.section}.{p.key}]')
                print(f'│     当前值: {p.current_value} → 建议值: {p.suggested_value}')
                print(f'│     理由: {p.reason}')
                print(f'│     置信度: {p.confidence:.0%}')
            print(f'└{"─"*76}┘')

        # ─── 摘要 ───
        print(f'\n{"="*78}')
        print(f'  复盘摘要')
        print(f'  {"─"*56}')
        print(f'  总股票数: {summary["total_stocks"]:<6d}  |  '
              f'上涨趋势: {summary["up_trend"]:<4d}  |  '
              f'盘整: {summary["consolidation"]:<4d}  |  '
              f'下跌趋势: {summary["down_trend"]}')
        print(f'  买点信号: {summary["buy_signals"]:<6d}  |  '
              f'卖点信号: {summary["sell_signals"]:<6d}  |  '
              f'A级机会: {summary["a_rank"]:<4d}  |  '
              f'B级机会: {summary["b_rank"]}')
        print(f'  双级别确认买: {summary["confirmed_buy"]:<4d}  |  '
              f'二买确认: {summary["second_buy"]:<4d}  |  '
              f'高危股票: {summary["high_risk"]}')
        print(f'{"="*78}\n')

    def _print_signal_section(
        self, groups: Dict[str, List[StockDiagnosis]],
        title: str, n: int, is_buy: bool
    ):
        """打印信号区域（机会或风险）"""
        section_config = {
            'A_strong_buy': ('🟢 A级 强买信号（多级别确认）', ['A_strong_buy', 'B2_confirmed', 'B1_reversal', 'B3_acceleration', 'double_confirmed']),
            'confirmed_sell': ('🔴 多级别确认卖点', ['confirmed_sell', 'S1_top_reversal', 'S3_breakdown', 'top_divergence', 'ma_breakdown']),
        }

        if is_buy:
            ordered_keys = ['A_strong_buy', 'B2_confirmed', 'B1_reversal', 'B3_acceleration', 'double_confirmed']
            labels = {
                'A_strong_buy': '🟢 A级 强买（多级别确认）',
                'B2_confirmed': '🔵 B2 二买确认（最安全买点）',
                'B1_reversal': '🟡 B1 一买（底部反转）',
                'B3_acceleration': '🟣 B3 三买（趋势加速）',
                'double_confirmed': '⭐ 双级别确认',
            }
        else:
            ordered_keys = ['confirmed_sell', 'S1_top_reversal', 'S3_breakdown', 'top_divergence', 'ma_breakdown']
            labels = {
                'confirmed_sell': '🔴 多级别确认卖点',
                'S1_top_reversal': '🟠 S1 一卖（顶部反转）',
                'S3_breakdown': '🟤 S3 三卖（趋势破位）',
                'top_divergence': '⚠ 顶背驰预警',
                'ma_breakdown': '📉 均线破位（≥2条）',
            }

        print(f'\n┌{"─"*76}┐')
        print(f'│  {title:<70s}│')
        print(f'├{"─"*76}┤')

        any_shown = False
        for key in ordered_keys:
            snaps = groups.get(key, [])
            if not snaps:
                continue
            any_shown = True
            label = labels.get(key, key)
            print(f'│  [{label}] 共 {len(snaps)} 只{" " * (62 - len(label) - len(str(len(snaps))))}│')
            if not is_buy and key == 'ma_breakdown':
                # 均线破位特殊表头
                pass
            else:
                print(f'│  {"代码":<10s} {"名称":<10s} {"行业":<14s} '
                      f'{"收盘":>7s} {"涨跌":>6s} {"走势":>4s} '
                      f'{"中枢位":>5s} {"置信":>5s} {"综合":>6s} │')

            trend_map = {2: '上涨', 1: '盘整', 0: '--', -2: '下跌'}
            pp_map = {1: '上方', 0: '内部', -1: '下方'}

            for s in snaps[:n]:
                conf = s.buy_confidence if is_buy else s.sell_confidence
                print(f'│  {s.symbol:<10s} {s.name:<10s} {s.industry:<14s} '
                      f'{s.close:>7.2f} {s.change_pct:>+5.1f}% '
                      f'{trend_map.get(s.trend_type, "?"):>4s} '
                      f'{pp_map.get(s.pivot_position, "?"):>5s} '
                      f'{conf:>4.0%} '
                      f'{s.composite_score:>+5.2f} │')

        if not any_shown:
            if is_buy:
                print(f'│  今日无显著买点信号{" " * 56}│')
            else:
                print(f'│  今日无显著风险信号{" " * 56}│')

        print(f'└{"─"*76}┘')

    # ========== 导出 ==========

    def _build_llm_data(self, report: DailyReviewReport) -> str:
        """构建传给LLM的市场分析原始数据 + 缠论课文提炼 + 参考知识"""
        m = report.market
        s = report.summary
        L = []

        # ---- 加载股票名称映射 ----
        stock_names = {}
        name_file = os.path.join(PROJECT_ROOT, 'data', 'stock_data', 'stock_metadata', 'stock_list.csv')
        if os.path.exists(name_file):
            try:
                import pandas as pd
                nm_df = pd.read_csv(name_file, dtype={'symbol': str})
                stock_names = dict(zip(nm_df['symbol'], nm_df['name']))
            except Exception:
                pass

        def _nm(code):
            """获取股票名称, 找不到返回空"""
            n = stock_names.get(str(code).zfill(6), '')
            return n

        L.append(f'复盘日期: {report.date}')
        L.append('')

        # ---- 大盘环境（完整数据） ----
        L.append('== 大盘环境体检 ==')
        L.append(f'上证指数收盘: {m.index_close:.1f}, 涨跌幅: {m.index_change_pct:+.2f}%')
        L.append(f'量能: {m.index_volume}(量比{m.volume_ratio:.1f}x), 波动率: {m.volatility_pct:.1%}')
        L.append(f'市场状态: {m.regime}(置信度{m.regime_confidence:.0%}), 情绪周期: {m.sentiment_cycle}, 风格: {m.style_regime}')
        L.append(f'动量评分: {m.momentum_score:+.2f}, 趋势评分: {getattr(m, "trend_score", 0):+.2f}')
        L.append(f'涨跌分布: {m.advancing}涨/{m.declining}跌(上涨率{m.up_pct:.0f}%)')
        L.append(f'涨停/跌停: {m.limit_up_count}家/{m.limit_down_count}家(比值{m.limit_up_ratio:.1f})')
        L.append(f'指数分化: {getattr(m, "index_divergence", False)}, 熊市风险: {m.bear_risk}')
        L.append(f'综合判定: {m.summary}')

        # ---- 量价原始数据（让LLM自己分析） ----
        chg_pct = m.index_change_pct
        vol_ratio = m.volume_ratio
        vol_label = m.index_volume
        adv_pct = m.up_pct
        L.append('== 量价原始数据 ==')
        L.append(f'上证涨跌幅: {chg_pct:+.2f}%')
        L.append(f'量比(今日量/20日均量): {vol_ratio:.1f}x, 量能标签: {vol_label}')
        L.append(f'涨跌比: {m.advancing}涨/{m.declining}跌 (上涨率{adv_pct:.0f}%)')
        L.append(f'涨停{m.limit_up_count}家/跌停{m.limit_down_count}家 (比值{m.limit_up_ratio:.1f})')
        L.append('量价关系参考知识: 放量+上涨=多头占优; 缩量+上涨=可能反弹非反转; 放量+下跌=空头占优; 缩量+下跌=卖盘枯竭; 上涨率>60%=普涨, 30-60%=分化, <30%=普跌')
        L.append(f'市场状态: {m.regime}, 情绪周期: {m.sentiment_cycle}, 风格: {m.style_regime}')
        L.append(f'动量: {m.momentum_score:+.2f}, 波动率: {m.volatility_pct:.1%}, 系统判定: {m.summary}, 熊市风险: {m.bear_risk}')

        # ---- 板块原始数据（含全部个股明细） ----
        L.append('')
        L.append('== 板块原始数据（含全部个股明细，引用时必须列出所有个股的表格） ==')
        top_inds = sorted(report.main_lines, key=lambda x: x.hot_rank)[:10]

        # 构建 industry → all stocks 映射
        ind_all_stocks = defaultdict(list)
        ind_buy_stocks_map = defaultdict(list)
        for d in report.all_diagnoses:
            if d.industry:
                ind_all_stocks[d.industry].append(d)
                if d.buy_point > 0:
                    ind_buy_stocks_map[d.industry].append(d)

        for ind in top_inds:
            buy_list = ind_buy_stocks_map.get(ind.industry, [])
            buy_list.sort(key=lambda x: x.composite_score, reverse=True)
            all_list = ind_all_stocks.get(ind.industry, [])
            # 板块涨跌统计
            up_count = sum(1 for d in all_list if d.change_pct > 0)
            down_count = sum(1 for d in all_list if d.change_pct < 0)
            L.append(f'#{ind.hot_rank} {ind.industry}: {ind.stock_count}只, '
                     f'上涨{up_count}只/下跌{down_count}只, '
                     f'上涨趋势{ind.up_trend_pct:.0f}%/下跌趋势{ind.down_trend_pct:.0f}%, '
                     f'B1:{ind.b1_count}/B2:{ind.b2_count}/B3:{ind.b3_count}, '
                     f'梯队{ind.formation_score:.1f}/3.0 | {ind.summary}')
            if buy_list:
                L.append(f'  【该板块全部买点个股({len(buy_list)}只)，必须全部列入表格】:')
                for d in buy_list:
                    bp = {1: 'B1', 2: 'B2', 3: 'B3'}.get(d.buy_point, '')
                    if d.second_buy: bp = 'B2*'
                    nm = d.name or _nm(d.symbol)
                    L.append(f'    {d.symbol} {nm} | {bp} | 收盘{d.close:.2f} | {d.change_pct:+.1f}% | 评分{d.composite_score:+.2f}')
        L.append('板块梯队参考: 完整=B3龙头+B2中军+B1后备都有(行情可持续); 残缺=单一买点(可能一日游)')
        L.append('重要: 分析每个行业时，必须用表格列出该板块【全部】买点个股，一只都不能少')

        # ---- 风格数据 ----
        growth_list = ['人工智能/算力', '互联网/软件', '半导体/光伏', '电子', '新能源车/风电', '通信/计算机', '医药', '传媒']
        value_list = ['金融', '基建/地产/石油石化', '有色/钢铁/煤炭', '消费', '交运', '建材', '化工']
        g_active = [i for i in top_inds if i.industry in growth_list and i.b1_count+i.b2_count+i.b3_count >= 3]
        v_active = [i for i in top_inds if i.industry in value_list and i.b1_count+i.b2_count+i.b3_count >= 3]
        L.append('')
        L.append(f'活跃成长板块({len(g_active)}个): {", ".join(i.industry for i in g_active) or "无"}')
        L.append(f'活跃价值板块({len(v_active)}个): {", ".join(i.industry for i in v_active) or "无"}')
        L.append('注: 成长板块多=风险偏好高; 价值板块多=偏防御')

        # ---- 涨停板 ----
        if report.limit_up_summary and report.limit_up_summary.limit_up_stocks:
            lu = report.limit_up_summary
            L.append('')
            L.append(f'涨停总数: {lu.total_limit_up}家, 主导模式: {lu.dominant_pattern}')
            if lu.common_traits:
                L.append(f'涨停共性: {"; ".join(lu.common_traits[:5])}')
            L.append(f'涨停总数: {lu.total_limit_up}家, 主导模式: {lu.dominant_pattern}')
            if lu.common_traits:
                L.append(f'涨停共性: {"; ".join(lu.common_traits[:5])}')
            L.append('【全部涨停个股明细，必须全部列入表格】:')
            for d in lu.limit_up_stocks:
                nm = d.name or _nm(d.symbol)
                L.append(f'{d.symbol} {nm}: {d.limit_up_type}(置信{d.type_confidence:.0%}), '
                         f'涨幅{d.change_pct:+.1f}%, 量比{d.vol_spike:.1f}x, '
                         f'前结构:{d.pre_trend_type}/{d.pre_pivot_position}, {d.chan_explanation[:100]}')

        # ---- 买点机会（预格式化表格，LLM直接复制） ----
        L.append('')
        total_buy = sum(1 for d in report.all_diagnoses if d.buy_point > 0)
        b1_total = sum(1 for d in report.all_diagnoses if d.buy_point == 1)
        b2_total = sum(1 for d in report.all_diagnoses if d.second_buy)
        b3_total = sum(1 for d in report.all_diagnoses if d.buy_point == 3)
        L.append(f'== 买点机会（全市场共{total_buy}只买点信号: B1={b1_total}只, B2={b2_total}只, B3={b3_total}只） ==')
        L.append('以下表格直接复制到报告中，不要修改或截断。')

        buy_descriptions = {
            'A_strong_buy': '笔线段双级别共振，可靠性最高',
            'B2_confirmed': '一买后回踩不破前低，最安全买点(第49课)',
            'B1_reversal': '下跌趋势末端底背驰，左侧抄底(第21课)',
            'B3_acceleration': '突破中枢上沿回踩确认，主升浪起点(第54课)',
        }
        op_labels = {'A_strong_buy': 'A级强买', 'B2_confirmed': 'B2二买最安全',
                     'B1_reversal': 'B1一买底部反转', 'B3_acceleration': 'B3三买趋势加速'}
        trend_map = {2: '上涨趋势', 1: '盘整', -2: '下跌趋势', 0: '-'}
        pp_map = {1: '中枢上', 0: '中枢内', -1: '中枢下'}
        bp_label_map = {1: 'B1', 2: 'B2', 3: 'B3'}
        SHOW_MAX = 10

        # 概览表
        L.append('【买点概览表，直接复制】')
        L.append('| 买点类型 | 数量 | 说明 |')
        L.append('|----------|------|------|')
        for key, label in op_labels.items():
            snaps = report.opportunities.get(key, [])
            if snaps:
                L.append(f'| {label} | {len(snaps)}只 | {buy_descriptions.get(key, "")} |')

        # 明细表
        for key, label in op_labels.items():
            snaps = report.opportunities.get(key, [])
            if not snaps:
                continue
            n_total = len(snaps)
            L.append('')
            if n_total <= SHOW_MAX:
                L.append(f'【{label} — 共{n_total}只，全部如下，直接复制此表格】')
                L.append(f'| 代码 | 名称 | 行业 | 收盘 | 涨跌 | 买点 | 走势 | 中枢位 | 置信 | 背驰 | 综合 | 止损 |')
                L.append(f'|------|------|------|------|------|------|------|--------|------|------|------|------|')
                for d in snaps:
                    bp = bp_label_map.get(d.buy_point, '')
                    if d.second_buy: bp = 'B2*'
                    nm = d.name or _nm(d.symbol)
                    ind = d.industry if d.industry and d.industry != '未归类' else '未归类'
                    div = '底背驰' if d.bottom_divergence else ('顶背驰' if d.top_divergence else '-')
                    L.append(f'| {d.symbol} | {nm} | {ind} | {d.close:.2f} | {d.change_pct:+.1f}% | {bp} | '
                             f'{trend_map.get(d.trend_type, "?")} | {pp_map.get(d.pivot_position, "?")} | '
                             f'{d.buy_confidence:.0%} | {div} | {d.composite_score:+.2f} | {d.stop_loss_price:.2f} |')
            else:
                L.append(f'【{label} — 共{n_total}只，前{SHOW_MAX}只明细表+全部代码表，直接复制以下两个表格】')
                L.append(f'（精选前{SHOW_MAX}只明细）')
                L.append(f'| 代码 | 名称 | 行业 | 收盘 | 涨跌 | 买点 | 走势 | 中枢位 | 置信 | 背驰 | 综合 | 止损 |')
                L.append(f'|------|------|------|------|------|------|------|--------|------|------|------|------|')
                for d in snaps[:SHOW_MAX]:
                    bp = bp_label_map.get(d.buy_point, '')
                    if d.second_buy: bp = 'B2*'
                    nm = d.name or _nm(d.symbol)
                    ind = d.industry if d.industry and d.industry != '未归类' else '未归类'
                    div = '底背驰' if d.bottom_divergence else ('顶背驰' if d.top_divergence else '-')
                    L.append(f'| {d.symbol} | {nm} | {ind} | {d.close:.2f} | {d.change_pct:+.1f}% | {bp} | '
                             f'{trend_map.get(d.trend_type, "?")} | {pp_map.get(d.pivot_position, "?")} | '
                             f'{d.buy_confidence:.0%} | {div} | {d.composite_score:+.2f} | {d.stop_loss_price:.2f} |')
                # 全部代码表
                all_codes = [d.symbol for d in snaps]
                L.append(f'（全部{n_total}只代码）')
                L.append(', '.join(all_codes))

        # ---- 风险预警（预格式化表格，LLM直接复制） ----
        L.append('')
        L.append('== 风险预警（以下表格直接复制到报告中，不要修改或截断） ==')
        risk_labels = {'confirmed_sell': '多级别确认卖点', 'S1_top_reversal': '一卖顶部反转',
                       'S3_breakdown': '三卖趋势破位', 'top_divergence': '顶背驰预警', 'ma_breakdown': '均线破位'}
        risk_descriptions = {
            'confirmed_sell': '上涨趋势末端，笔线段双级别共振卖点',
            'S1_top_reversal': '上涨趋势+中枢上方，顶背驰确认',
            'S3_breakdown': '跌破中枢下沿，反抽不涨回',
            'top_divergence': 'MACD顶背离预警',
            'ma_breakdown': '股价跌破关键均线支撑（≥2条）',
        }
        trend_map2 = {2: '上涨趋势', 1: '盘整', -2: '下跌趋势', 0: '-'}
        pp_map2 = {1: '中枢上', 0: '中枢内', -1: '中枢下'}
        sp_map = {1: 'S1', 2: 'S2', 3: 'S3'}

        # 概览表
        L.append('【风险概览表，直接复制】')
        L.append('| 风险类型 | 数量 | 说明 |')
        L.append('|----------|------|------|')
        for key, label in risk_labels.items():
            snaps = report.risk_alerts.get(key, [])
            if snaps:
                L.append(f'| {label} | {len(snaps)}只 | {risk_descriptions.get(key, "")} |')

        # 明细表
        SHOW_MAX = 10
        for key, label in risk_labels.items():
            snaps = report.risk_alerts.get(key, [])
            if not snaps:
                continue
            n_total = len(snaps)
            L.append('')
            if n_total <= SHOW_MAX:
                L.append(f'【{label} — 共{n_total}只，全部如下，直接复制此表格】')
                L.append(f'| 代码 | 名称 | 行业 | 收盘 | 涨跌 | 卖点 | 走势 | 中枢位 | 置信 | 综合 |')
                L.append(f'|------|------|------|------|------|------|------|--------|------|------|')
                for d in snaps:
                    ind = d.industry if d.industry and d.industry != '未归类' else '未归类'
                    nm = d.name or _nm(d.symbol)
                    sp = sp_map.get(d.sell_point, '')
                    conf = d.sell_confidence if hasattr(d, 'sell_confidence') and d.sell_confidence > 0 else ''
                    L.append(f'| {d.symbol} | {nm} | {ind} | {d.close:.2f} | {d.change_pct:+.1f}% | {sp} | '
                             f'{trend_map2.get(d.trend_type, "?")} | {pp_map2.get(d.pivot_position, "?")} | '
                             f'{conf:.0%} | {d.composite_score:+.2f} |')
            else:
                L.append(f'【{label} — 共{n_total}只，前{SHOW_MAX}只明细表+全部代码表，直接复制以下两个表格】')
                # 精选表
                L.append(f'（精选前{SHOW_MAX}只明细）')
                L.append(f'| 代码 | 名称 | 行业 | 收盘 | 涨跌 | 卖点 | 走势 | 中枢位 | 置信 | 综合 |')
                L.append(f'|------|------|------|------|------|------|------|--------|------|------|')
                for d in snaps[:SHOW_MAX]:
                    ind = d.industry if d.industry and d.industry != '未归类' else '未归类'
                    nm = d.name or _nm(d.symbol)
                    sp = sp_map.get(d.sell_point, '')
                    conf = d.sell_confidence if hasattr(d, 'sell_confidence') and d.sell_confidence > 0 else ''
                    L.append(f'| {d.symbol} | {nm} | {ind} | {d.close:.2f} | {d.change_pct:+.1f}% | {sp} | '
                             f'{trend_map2.get(d.trend_type, "?")} | {pp_map2.get(d.pivot_position, "?")} | '
                             f'{conf:.0%} | {d.composite_score:+.2f} |')
                # 全部代码表
                all_codes = [d.symbol for d in snaps]
                L.append(f'（全部{n_total}只代码）')
                L.append(', '.join(all_codes))

        # ---- 全市场摘要 ----
        L.append('')
        L.append('== 全市场扫描摘要 ==')
        L.append(f'诊断股票: {s["total_stocks"]}只, 上涨趋势{s["up_trend"]}/盘整{s["consolidation"]}/下跌{s["down_trend"]}')
        L.append(f'买点信号: {s["buy_signals"]}只, 卖点信号: {s["sell_signals"]}只')
        L.append(f'A级{s["a_rank"]}只, B级{s["b_rank"]}只, 双级别确认{s["confirmed_buy"]}只, 高危{s["high_risk"]}只')
        L.append(f'涨停: {report.limit_up_summary.total_limit_up if report.limit_up_summary else 0}只')

        # ---- 缠论课文内容提炼 ----
        L.append('')
        L.append('== 缠论核心课文内容速查（引用时必须提炼内容，绝对不能只报课号） ==')
        lessons = [
            ('第11课《不会吻,无以高潮》',
             '均线系统的"吻": 飞吻(短期均线略碰就分开,趋势弱), 唇吻(触碰后继续原方向,趋势中继), '
             '湿吻(反复缠绕,趋势转折)。多头排列(短均线在长均线之上)=强势, 空头排列=弱势。'),
            ('第15课《没有趋势,没有背驰》',
             '背驰必须在趋势中判断 -- 没有趋势就没有背驰。盘整中的MACD变化不是背驰只是震荡。'
             '判断背驰前提: 先确定走势类型(上涨或下跌趋势), 再比较同方向两段力度, 力度衰竭=背驰成立。'),
            ('第17课《走势终完美》',
             '"走势终完美"是缠论第一原理。任何级别的任何走势类型终将完成 -- 上涨必然结束, 下跌也必然结束。'
             '走势完成后必然转化为其他走势类型。这个原理保证了买卖点的必然存在: 下跌结束必然出现买点。'),
            ('第18课《不被面首的萌动是安全的》',
             '"中枢"的定义: 某级别走势类型中被至少三个连续次级别走势类型重叠的部分。中枢是多空反复争夺的'
             '价格区间(可理解为"横盘震荡区")。中枢之上=强势, 中枢之下=弱势, 中枢内部=方向不明。'),
            ('第21课《缠中说禅买卖点分析的完备性》',
             '最重要的课文之一。三类买卖点的完备定义: 一买=下跌趋势的底背驰终点(抄底); '
             '二买=一买后回踩不创新低(加仓最安全); 三买=突破中枢上沿后回踩不跌回中枢(追涨确认)。'
             '对应三卖: 一卖=上涨趋势顶背驰(逃顶), 二卖=反弹不创新高, 三卖=跌破中枢后反抽不涨回。'
             '关键定理: 任何买卖点都必然对应着某级别的中枢。'),
            ('第24课《MACD对背驰的辅助判断》',
             'MACD是判断背驰的辅助工具。股价新高/新低但MACD柱面积缩小, 黄白线未能同步新高/新低时, '
             '背驰成立。柱面积比较: 后一段面积明显小于前一段=背驰。注意MACD只是辅助, 最终要结合走势结构。'),
            ('第35课《给基础差的同学补补课》',
             '多级别联立分析: 大级别定方向, 小级别定买卖点。日线定中期趋势, 60分钟/30分钟找精确买卖点。'
             '不同级别同时发信号=多级别共振, 可靠性远高于单一级别。这是"双级别确认"的理论基础。'),
            ('第49课《利润率最大的操作模式》',
             '第二类买点是最安全, 利润率最大的买点。它既有一买作为"底"的保障, 又不需要像一买那样承担左侧抄底风险。'
             '操作: 等待一买出现后, 股价回踩不破一买低点, MACD再次金叉时介入。'),
            ('第54课《一个具体走势的分析》',
             '三买的确认: 股价突破最近一个中枢的上沿后, 次级别回踩不跌回中枢区间。三买标志中枢被有效突破, '
             '价格进入新的上涨阶段, 是主升浪的起点。需成交量配合, 放量突破更可靠。'),
            ('第89课《中阴阶段的具体分析》',
             '"中阴"是旧趋势结束, 新趋势未形成的过渡阶段。操作原则: 看不懂的就不做 -- '
             '方向不明时持仓等待或减仓, 不要频繁交易。中阴结束标志: 出现明确的三买或三卖。'),
        ]
        for title, content in lessons:
            L.append(f'{title}: {content}')

        # ---- 当前持仓评估 ----
        holdings = report.portfolio_holdings
        if holdings:
            L.append('')
            total_mv = sum(h['market_value'] for h in holdings)
            total_pnl = total_mv - sum(h['cost']*h['shares'] for h in holdings)
            L.append(f'== 当前持仓评估(资金: {report.portfolio_cash:,.0f}元, {len(holdings)}只, '
                     f'持仓市值{total_mv:,.0f}, 浮动盈亏{total_pnl:+,.0f}) ==')
            L.append('')
            L.append('【持仓汇总表】')
            L.append('| 代码 | 名称 | 持仓(股) | 成本 | 现价 | 市值 | 盈亏% | 走势 | 买点 | 卖点 | 风险 | 建议 |')
            L.append('|------|------|----------|------|------|------|-------|------|------|------|------|------|')
            for h in holdings:
                L.append(f'| {h["code"]} | {h["name"]} | {h["shares"]:,.0f} | {h["cost"]:.2f} | '
                         f'{h["close"]:.2f} | {h["market_value"]:,.0f} | {h["pnl_pct"]:+.1f}% | '
                         f'{h["trend"]} | {h["buy_pt"]} | {h["sell_pt"]} | '
                         f'{h["risk"]} | {h["action"]} |')
            L.append('')
            L.append('以下为每只持仓个股的深度诊断数据，必须逐只写分析段落：')
            for h in holdings:
                L.append(f'--- {h["code"]} {h["name"]} ---')
                for line in h['analysis'].split('\n'):
                    L.append(f'  {line}')
            L.append('')
            L.append('【持仓评估写作要求】')
            L.append('1. 先列出持仓汇总表（使用上面的表格数据）')
            L.append('2. 然后逐只写深度分析段落，每只至少5行：')
            L.append('   - 缠论结构定位：现在处于什么走势、中枢位置、笔/线段方向')
            L.append('   - 买卖点评估：是否有买点支撑或卖点威胁，背驰情况')
            L.append('   - 量价关系判断：放量/缩量状态是否健康')
            L.append('   - 关键价位操作计划：止损价、止盈目标、加仓/减仓触发条件')
            L.append('   - 风险提示：均线破位、顶背驰、卖点等风险信号')
            L.append('3. 最后给出整体持仓评估和仓位建议')
        else:
            L.append('')
            L.append('== 当前持仓: 空仓 ==')
            L.append('（如无持仓，报告中简要说明"当前空仓"并直接跳到操作计划部分）')

        # ---- 持仓替换计划 ----
        replacements = report.replacement_plan
        if replacements:
            L.append('')
            L.append(f'== 持仓替换计划(共{len(replacements)}条) ==')
            L.append('| 持仓股 | 操作 | 原因 | 替换标的 | 评分 | 现价 | 紧急度 |')
            L.append('|--------|------|------|----------|------|------|--------|')
            for r in replacements:
                L.append(f'| {r["code"]} {r["name"]} | {r["action"]} | {r["reason"]} | '
                         f'{r["replacement_code"]} {r["replacement_name"]} | '
                         f'{r["replacement_score"]:.3f} | {r["replacement_price"]:.2f} | '
                         f'{r["urgency"]} |')
            L.append('')
            L.append('【替换指令】"六、持仓评估与操作计划"中必须列出上述替换计划。')
            L.append('立即替换的标的需要在操作计划表中写明具体买入/卖出方案。')

        # ---- 当日选股结果 ----
        selections = report.today_selections
        if selections:
            L.append('')
            L.append(f'== 当日策略选股结果(共{len(selections)}只) — 这是"四、机会挖掘"深度分析的主角 ==')
            L.append('| # | 代码 | 名称 | 行业 | 评分 | 权重 | 目标仓位 | 现价 | 信号 | 选股理由 |')
            L.append('|---|------|------|------|------|------|----------|------|------|----------|')
            for i, s in enumerate(selections):
                code = str(s.get('股票代码', ''))
                name = str(s.get('股票名称', ''))
                ind = str(s.get('行业', ''))
                score = s.get('综合评分', 0)
                weight = s.get('权重', 0)
                target = s.get('目标仓位', 0)
                price = s.get('当前价格', 0)
                signal = str(s.get('信号强度', ''))
                reason = str(s.get('选股理由', ''))[:80]
                L.append(f'| {i+1} | {code} | {name} | {ind} | {float(score):.3f} | '
                         f'{float(weight)*100:.1f}% | {float(target):,.0f} | {float(price):.2f} | {signal} | {reason} |')
            L.append('')
            L.append('【重要指令】"四、机会挖掘"的"重点个股深度分析"必须以上述策略选股结果为对象，')
            L.append('逐只写深度分析段落。不要从全市场买点信号中自行挑选。')
            L.append('每只策略选股写一段分析: 缠论结构定位 + 买卖点评估 + 量价判断 + 操作建议。')

        # ---- 系统进化建议 ----
        if report.evolution_patches:
            L.append('')
            L.append(f'== 系统进化建议(共{len(report.evolution_patches)}条) ==')
            for p in report.evolution_patches:
                L.append(f'- [{p.section}.{p.key}] {p.current_value}→{p.suggested_value} ({p.urgency})')
                L.append(f'  理由: {p.reason}')

        # ---- 准确率历史 ----
        ah = report.accuracy_history
        if ah:
            L.append('')
            L.append(f'== 系统准确率进化记录(共{len(ah)}条) ==')
            L.append('| 日期 | 方向准确率 | 平均收益 | 超额收益 | 信号数 |')
            L.append('|------|------------|----------|----------|--------|')
            for h in ah:
                L.append(f'| {h["date"]} | {h["direction_accuracy"]:.0%} | '
                         f'{h["avg_return"]:+.2f}% | {h["excess_return"]:+.2f}% | '
                         f'{h["hits"]}/{h["total_signals"]} |')
            # 趋势
            if len(ah) >= 3:
                recent_accs = [h['direction_accuracy'] for h in ah[-5:]]
                avg_recent = sum(recent_accs) / len(recent_accs)
                L.append(f'近5次平均准确率: {avg_recent:.0%}')

        return '\n'.join(L)

    def _export_llm_report(self, report: DailyReviewReport, date_str: str,
                           chart_files: list = None) -> str:
        """使用DeepSeek生成叙事风格复盘报告"""
        api_key = self._get_llm_api_key()
        if not api_key:
            print('  [WARN] DeepSeek API Key未配置, 跳过LLM报告')
            return ''

        # 从外部文件加载系统提示
        prompt_file = os.path.join(STRATEGY_DIR, 'config', 'review_prompt.txt')
        if os.path.exists(prompt_file):
            with open(prompt_file, 'r', encoding='utf-8') as f:
                system_prompt = f.read()
        else:
            print(f'  [WARN] 未找到提示词文件: {prompt_file}')
            return ''

        review_data = self._build_llm_data(report)

        # K线图引用
        chart_info = ''
        if chart_files:
            chart_info = '\n\n已生成真实K线+MACD图表(png文件), 报告中请引用:\n'
            for cf in chart_files:
                chart_info += f'  [{cf["symbol"]} {cf["name"]}] K线+MACD图: {cf["rel_path"]}\n'
            chart_info += '请在个股分析处标注"[K线图: charts/XXXXXX_Kline_MACD.png]", 读者可打开对应图片查看真实K线走势和MACD背离。\n'

        user_prompt = f'请根据以下复盘数据和预格式化表格, 生成{report.date}的每日复盘文章。\n\n{chart_info}\n注意：数据中已包含预格式化的Markdown表格（以【标记），这些表格必须原样复制到报告中，不要修改或截断。\n\n{review_data}\n\n请按照风格要求, 将这些数据转化为一篇专业的复盘文章。所有【标记的表格必须完整保留。'

        try:
            import requests
            resp = requests.post(
                'https://api.deepseek.com/v1/chat/completions',
                headers={'Authorization': f'Bearer {api_key}', 'Content-Type': 'application/json'},
                json={
                    'model': 'deepseek-chat',
                    'messages': [
                        {'role': 'system', 'content': system_prompt},
                        {'role': 'user', 'content': user_prompt},
                    ],
                    'temperature': 0.7,
                    'max_tokens': 65536,
                },
                timeout=300,
            )
            resp.raise_for_status()
            content = resp.json()['choices'][0]['message']['content']
        except Exception as e:
            print(f'  [WARN] LLM报告生成失败: {e}')
            return ''

        # 生成标题（LLM经常漏掉或生成劣质标题，Python直接生成并覆盖）
        sentiment_kw = {'冰点': '冰点情绪', '修复': '情绪修复', '高潮': '情绪高潮',
                        '分歧': '市场分歧', '退潮': '情绪退潮'}
        regime_kw = {'牛市': '牛市格局', '震荡': '震荡格局', '熊市': '熊市防御'}
        sk = sentiment_kw.get(report.market.sentiment_cycle, '市场')
        rk = regime_kw.get(report.market.regime, '')
        b3_n = sum(1 for d in report.all_diagnoses if d.buy_point == 3)
        main_lines = [ind for ind in report.main_lines if ind.summary.startswith('★')]
        top = main_lines[0].industry if main_lines else '主线'
        theme = f'聚焦{top}B3机会' if b3_n > 20 else ('等待结构信号' if report.market.sentiment_cycle == '冰点' else '结构分化中挖掘机会')
        title_line = f'# {report.date} 每日复盘：{sk}下的{rk}，{theme}'

        # 去掉LLM可能自己生成的标题（以 # 开头的第一行）
        content = content.strip()
        if content.startswith('# '):
            first_newline = content.find('\n')
            if first_newline > 0:
                content = content[first_newline:].strip()
        content = title_line + '\n\n' + content

        output_path = f'复盘报告_公众号版_{report.date}.md'
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(content)

        # 同时生成PDF(内容一致)
        try:
            pdf_path = self._md_to_pdf(content, report.date)
            if pdf_path:
                print(f'  LLM报告PDF: {pdf_path}')
        except Exception as e:
            print(f'  [WARN] PDF生成失败: {e}')

        return output_path

    def _md_to_pdf(self, markdown_text: str, date_label: str) -> str:
        """将Markdown内容转为PDF, 嵌入K线图"""
        try:
            from fpdf import FPDF
        except ImportError:
            return ''

        font_paths = [
            '/mnt/c/Windows/Fonts/simhei.ttf',
            '/mnt/c/Windows/Fonts/simsun.ttf',
            '/mnt/c/Windows/Fonts/msyh.ttc',
        ]
        font_path = ''
        for fp in font_paths:
            if os.path.exists(fp):
                font_path = fp
                break
        if not font_path:
            return ''

        pdf = FPDF()
        pdf.set_auto_page_break(auto=True, margin=12)
        # 更窄的左右边距，避免内容飞出页面
        pdf.set_left_margin(10)
        pdf.set_right_margin(10)
        pdf.add_font('CJK', '', font_path, uni=True)
        pdf.add_font('CJK', 'B', font_path, uni=True)
        pdf.add_page()

        # 可用宽度: A4=210mm, 左右margin=10mm → 190mm
        TEXT_W = 190
        # 表格可用宽度略窄（留出边距）
        TABLE_W = 186

        in_code_block = False
        code_lines = []
        chart_dir = os.path.join(self.output_dir, 'charts')

        def flush_code():
            nonlocal code_lines
            if code_lines:
                pdf.set_font('CJK', '', 6.5)
                pdf.set_fill_color(245, 245, 245)
                for cl in code_lines:
                    if pdf.get_y() > 272:
                        pdf.add_page()
                    pdf.set_x(10)
                    # 用 multi_cell 避免长行飞出页面
                    pdf.multi_cell(TEXT_W, 3.8, cl[:150], fill=True)
                code_lines = []
                pdf.ln(1)

        def try_embed_chart(line_text):
            """检测并嵌入K线图"""
            import re
            m = re.search(r'charts/(\d+)_Kline_MACD\.png', line_text)
            if m:
                img_path = os.path.join(chart_dir, f'{m.group(1)}_Kline_MACD.png')
                if os.path.exists(img_path):
                    if pdf.get_y() > 200:
                        pdf.add_page()
                    try:
                        pdf.image(img_path, x=10, w=TABLE_W)
                        pdf.set_font('CJK', '', 6)
                        pdf.set_text_color(120, 120, 120)
                        pdf.cell(0, 4, f'[图: {m.group(1)} K线+MACD]', align='C', new_x='LMARGIN', new_y='NEXT')
                        pdf.set_text_color(0, 0, 0)
                        pdf.ln(3)
                        return True
                    except Exception:
                        pass
            return False

        def render_table_row(cells, is_header=False):
            """渲染一行表格，使用 multi_cell 确保内容不溢出"""
            n = len(cells)
            if n == 0:
                return
            # 每列最小宽度10mm，最大宽度根据列数动态调整
            col_w = max(10, TABLE_W / n)
            # 如果总宽度超过可用宽度，压缩到恰好填满
            if col_w * n > TABLE_W:
                col_w = TABLE_W / n

            # 先计算本行最大高度
            pdf.set_font('CJK', 'B' if is_header else '', 5.5)
            cell_h = 4.0
            max_lines = 1
            for c in cells:
                # 估算需要多少行：中文每字约3mm宽(@5.5pt)，英文每字约1.5mm
                text_w = sum(3.0 if ord(ch) > 127 else 1.5 for ch in c)
                lines = max(1, int(text_w / (col_w - 1) + 0.9))
                max_lines = max(max_lines, min(lines, 4))  # 最多4行
            row_h = cell_h * max_lines

            # 检查是否需要换页
            if pdf.get_y() + row_h > 275:
                pdf.add_page()

            # 画每个单元格
            x_start = pdf.get_x()
            y_start = pdf.get_y()
            max_y = y_start

            for i, c in enumerate(cells):
                x_pos = x_start + i * col_w
                pdf.set_xy(x_pos, y_start)
                # 画背景
                if is_header:
                    pdf.set_fill_color(31, 78, 121)
                    pdf.set_text_color(255, 255, 255)
                else:
                    pdf.set_fill_color(255, 255, 255)
                    pdf.set_text_color(0, 0, 0)
                # 用 rect 画边框和背景
                pdf.rect(x_pos, y_start, col_w, row_h, style='DF')
                # 写文字（居中多行）
                pdf.set_xy(x_pos + 0.5, y_start + 0.3)
                pdf.multi_cell(col_w - 1, cell_h, c, align='C')

            pdf.set_xy(x_start, y_start + row_h)
            if is_header:
                pdf.set_text_color(0, 0, 0)

        for line in markdown_text.split('\n'):
            # 代码块
            if line.strip().startswith('```'):
                if in_code_block:
                    flush_code()
                    in_code_block = False
                else:
                    in_code_block = True
                continue
            if in_code_block:
                code_lines.append(line)
                continue

            # 嵌入K线图
            if try_embed_chart(line):
                continue

            # 空行
            if not line.strip():
                pdf.ln(2)
                continue

            # 标题 — 全部改用 multi_cell 防溢出
            if line.startswith('#### '):
                flush_code()
                pdf.set_font('CJK', 'B', 10)
                pdf.multi_cell(TEXT_W, 6, line[5:].strip())
                pdf.ln(1)
            elif line.startswith('### '):
                flush_code()
                pdf.set_font('CJK', 'B', 11)
                pdf.multi_cell(TEXT_W, 7, line[4:].strip())
                pdf.ln(1)
            elif line.startswith('## '):
                flush_code()
                pdf.set_fill_color(31, 78, 121)
                pdf.set_text_color(255, 255, 255)
                pdf.set_font('CJK', 'B', 12)
                pdf.cell(TEXT_W, 8, f'  {line[3:].strip()}', fill=True, new_x='LMARGIN', new_y='NEXT')
                pdf.set_text_color(0, 0, 0)
                pdf.ln(2)
            elif line.startswith('# '):
                flush_code()
                pdf.set_font('CJK', 'B', 16)
                pdf.multi_cell(TEXT_W, 10, line[2:].strip(), align='C')
                pdf.ln(3)

            # 表格行 — 改进的渲染
            elif line.strip().startswith('|') and line.strip().endswith('|'):
                flush_code()
                cells = [c.strip() for c in line.strip().split('|')[1:-1]]
                # 跳过分隔行
                if all(c.startswith('---') or c.startswith('===') or c == '' or set(c) <= set('-:= ')
                       for c in cells):
                    continue
                # 判断是否为表头（通过内容特征：包含"代码/名称/行业/排名"等字段）
                is_header = any(
                    kw in str(cells) for kw in ['代码', '名称', '排名', '指标', '行业', '日期', '操作', '数值']
                )
                render_table_row(cells, is_header=is_header)

            # 列表项
            elif line.strip().startswith('- ') or line.strip().startswith('* '):
                flush_code()
                pdf.set_font('CJK', '', 7.5)
                pdf.set_x(10)
                pdf.cell(4, 4.5, '-')
                pdf.multi_cell(TEXT_W - 4, 4.5, line.strip()[2:])

            # 普通文本
            else:
                flush_code()
                pdf.set_font('CJK', '', 8)
                pdf.set_x(10)
                pdf.multi_cell(TEXT_W, 5, line.strip())

        flush_code()

        output_pdf = f'复盘报告_公众号版_{date_label}.pdf'
        try:
            pdf.output(output_pdf)
            return output_pdf
        except Exception as e:
            print(f'  [WARN] PDF生成失败: {e}')
            return ''

    def _get_llm_api_key(self) -> str:
        import os
        key = os.environ.get('DEEPSEEK_API_KEY', '')
        if key: return key
        try:
            from core.config_loader import load_config
            cfg = load_config()
            key = cfg.get('industry_sentiment.deepseek_api_key', '')
        except Exception: pass
        return key

    def _generate_charts(self, report: DailyReviewReport, date_str: str):
        """用mplfinance生成真实的K线+MACD图表"""
        try:
            import mplfinance as mpf
            import matplotlib.pyplot as plt
            import matplotlib
            matplotlib.use('Agg')
        except ImportError:
            print('  [WARN] mplfinance未安装, 跳过K线图生成 (pip install mplfinance)')
            return []

        # 中文字体
        try:
            matplotlib.rcParams['font.sans-serif'] = ['SimHei', 'Microsoft YaHei', 'WenQuanYi Micro Hei']
            matplotlib.rcParams['axes.unicode_minus'] = False
        except Exception:
            pass

        charts_dir = os.path.join(self.output_dir, 'charts')
        os.makedirs(charts_dir, exist_ok=True)
        chart_files = []

        # Top 5买点机会
        top_stocks = []
        for key in ['A_strong_buy', 'B2_confirmed', 'B1_reversal', 'B3_acceleration']:
            snaps = report.opportunities.get(key, [])
            for d in snaps:
                if d.buy_point > 0 and d.symbol not in {s.symbol for s in top_stocks}:
                    top_stocks.append(d)
                    if len(top_stocks) >= 5:
                        break
            if len(top_stocks) >= 5:
                break

        for i, d in enumerate(top_stocks):
            try:
                fp = os.path.join(self.data_dir, f'{d.symbol}_qfq.csv')
                if not os.path.exists(fp):
                    continue
                df = pd.read_csv(fp, parse_dates=['datetime'])
                df = df.set_index('datetime')
                # 取最近120个交易日
                df = df.tail(120)
                if len(df) < 30:
                    continue

                # 计算MACD
                ema12 = df['close'].ewm(span=12).mean()
                ema26 = df['close'].ewm(span=26).mean()
                macd_line = ema12 - ema26
                signal_line = macd_line.ewm(span=9).mean()
                macd_hist = macd_line - signal_line

                buy_label = {1: 'B1 Buy(1st)', 2: 'B2 Buy(2nd)', 3: 'B3 Buy(3rd)'}.get(d.buy_point, 'Buy')

                # 创建MACD子图
                apds = [
                    mpf.make_addplot(macd_line, panel=2, color='#1F4E79', width=1, label='MACD'),
                    mpf.make_addplot(signal_line, panel=2, color='#C05050', width=1, label='Signal'),
                    mpf.make_addplot(macd_hist, type='bar', panel=2, color='#808080',
                                    alpha=0.5, label='Histogram'),
                ]

                # 标题
                title = f'{d.symbol} {d.name} | {d.industry} | {buy_label} | Score:{d.composite_score:+.2f} Stop:{d.stop_loss_price:.2f}'

                fig, axes = mpf.plot(
                    df, type='candle', style='charles',
                    volume=True, addplot=apds,
                    title=title,
                    returnfig=True,
                    figsize=(12, 7),
                    panel_ratios=(3, 1, 2),
                )

                # 标注中枢区间
                if not np.isnan(d.pivot_zg) and not np.isnan(d.pivot_zd):
                    ax0 = axes[0]
                    ax0.axhline(y=d.pivot_zg, color='#1F4E79', linestyle='--', linewidth=0.8, alpha=0.6)
                    ax0.axhline(y=d.pivot_zd, color='#1F4E79', linestyle='--', linewidth=0.8, alpha=0.6)
                    ax0.fill_between([df.index[0], df.index[-1]], d.pivot_zd, d.pivot_zg,
                                    color='#1F4E79', alpha=0.05)
                    ax0.text(df.index[-1], d.pivot_zg, f' Pivot:{d.pivot_zg:.2f}',
                            fontsize=7, color='#1F4E79', va='bottom')
                    ax0.text(df.index[-1], d.pivot_zd, f' Pivot:{d.pivot_zd:.2f}',
                            fontsize=7, color='#1F4E79', va='top')

                # 标注止损线
                ax0 = axes[0]
                ax0.axhline(y=d.stop_loss_price, color='#C05050', linestyle=':', linewidth=0.8)
                ax0.text(df.index[-1], d.stop_loss_price, f' Stop:{d.stop_loss_price:.2f}',
                        fontsize=7, color='#C05050', va='bottom')

                # 标注买点(在最近的K线位置)
                last_idx = len(df) - 1
                ax0.annotate(buy_label,
                            xy=(last_idx, float(df['low'].iloc[-1]) * 0.97),
                            fontsize=10, color='red', fontweight='bold',
                            ha='center',
                            arrowprops=dict(arrowstyle='->', color='red', lw=1.5))

                chart_path = os.path.join(charts_dir, f'{d.symbol}_Kline_MACD.png')
                fig.savefig(chart_path, dpi=150, bbox_inches='tight')
                plt.close(fig)
                chart_files.append({
                    'symbol': d.symbol,
                    'name': d.name,
                    'industry': d.industry,
                    'buy_point': d.buy_point,
                    'path': chart_path,
                    'rel_path': f'charts/{d.symbol}_Kline_MACD.png',
                })
            except Exception as e:
                print(f'  [WARN] {d.symbol} K线图生成失败: {e}')
                continue

        if chart_files:
            print(f'  生成K线图: {len(chart_files)} 张 -> {charts_dir}/')
        return chart_files

    def _export_report(self, report: DailyReviewReport):
        """导出报告为CSV + JSON"""
        date_str = report.date.replace('-', '')

        # 市场体检 → JSON
        market_path = os.path.join(self.output_dir, f'market_vitals_{date_str}.json')
        with open(market_path, 'w') as f:
            json.dump({
                'date': report.date,
                'index_close': report.market.index_close,
                'index_change_pct': report.market.index_change_pct,
                'index_volume': report.market.index_volume,
                'regime': report.market.regime,
                'sentiment_cycle': report.market.sentiment_cycle,
                'advancing': report.market.advancing,
                'declining': report.market.declining,
                'limit_up': report.market.limit_up_count,
                'limit_down': report.market.limit_down_count,
                'bear_risk': report.market.bear_risk,
                'summary': report.market.summary,
            }, f, ensure_ascii=False, indent=2)

        # 行业主线 → CSV
        ind_rows = []
        for ind in report.main_lines:
            ind_rows.append({
                'industry': ind.industry, 'stock_count': ind.stock_count,
                'up_trend_pct': ind.up_trend_pct, 'down_trend_pct': ind.down_trend_pct,
                'b1': ind.b1_count, 'b2': ind.b2_count, 'b3': ind.b3_count,
                'formation_score': ind.formation_score,
                'avg_composite': ind.avg_composite_score,
                'summary': ind.summary, 'hot_rank': ind.hot_rank,
            })
        if ind_rows:
            pd.DataFrame(ind_rows).to_csv(
                os.path.join(self.output_dir, f'industry_lines_{date_str}.csv'), index=False
            )

        # 进化补丁 → JSON
        if report.evolution_patches:
            patch_path = os.path.join(self.output_dir, f'evolution_patch_{date_str}.json')
            with open(patch_path, 'w') as f:
                json.dump([{
                    'section': p.section, 'key': p.key,
                    'current': p.current_value, 'suggested': p.suggested_value,
                    'reason': p.reason, 'confidence': p.confidence, 'urgency': p.urgency,
                } for p in report.evolution_patches], f, ensure_ascii=False, indent=2)

        # 完整报告文档 → Markdown
        report_path = os.path.join(self.output_dir, f'daily_review_{date_str}.md')
        self._export_document(report, report_path)

        # 全量个股缠论分析表 → CSV
        chan_csv_path = self._export_full_chan_table(report, date_str)

        # Word文档导出（结构化报告）
        word_path = self._export_word(report_path, date_str)

        # Word文档导出（公众号版，LLM报告路径）
        llm_md_path = os.path.join(PROJECT_ROOT, 'strategy', f'复盘报告_公众号版_{report.date}.md')
        if os.path.exists(llm_md_path):
            llm_word_path = self._export_word(llm_md_path, date_str, suffix='_公众号版')

        print(f'\n报告已导出至: {self.output_dir}/')
        print(f'  完整报告: daily_review_{date_str}.md')
        if word_path:
            print(f'  Word文档: {os.path.basename(word_path)}')
        print(f'  市场体检: market_vitals_{date_str}.json')
        print(f'  行业主线: industry_lines_{date_str}.csv')
        print(f'  全量缠论表: {os.path.basename(chan_csv_path)}' if chan_csv_path else '')
        if report.evolution_patches:
            print(f'  进化补丁: evolution_patch_{date_str}.json')

    def _export_full_chan_table(self, report: DailyReviewReport, date_str: str) -> str:
        """导出全量个股缠论分析表 — CSV格式，包含所有诊断过的股票"""
        diagnoses = report.all_diagnoses
        if not diagnoses:
            print('  [WARN] 无个股诊断数据，跳过缠论表导出')
            return ''

        trend_map = {2: '上涨趋势', 1: '盘整', 0: '-', -2: '下跌趋势'}
        pp_map = {1: '中枢上', 0: '中枢内', -1: '中枢下'}
        buy_map = {0: '-', 1: 'B1(一买)', 2: 'B2(二买)', 3: 'B3(三买)'}
        sell_map = {0: '-', 1: 'S1(一卖)', 2: 'S2(二卖)', 3: 'S3(三卖)'}
        stroke_map = {1: '↑向上', -1: '↓向下', 0: '-'}
        seg_map = {1: '↑向上', -1: '↓向下', 0: '-'}
        lvl_map = {3: '双级确认', 2: '线段级', 1: '笔级', 0: '-', -3: '双级卖点', -2: '线段卖', -1: '笔卖'}

        rows = []
        for d in sorted(diagnoses, key=lambda x: x.composite_score, reverse=True):
            rows.append({
                '股票代码': d.symbol,
                '股票名称': d.name,
                '行业': d.industry,
                '收盘价': round(d.close, 2),
                '涨跌幅%': round(d.change_pct, 2),
                '量比': round(d.volume_ratio, 2),
                '走势类型': trend_map.get(d.trend_type, str(d.trend_type)),
                '趋势强度': round(d.trend_strength, 3),
                '中枢位置': pp_map.get(d.pivot_position, str(d.pivot_position)),
                '中枢上沿ZG': round(d.pivot_zg, 2) if not np.isnan(d.pivot_zg) else '',
                '中枢下沿ZD': round(d.pivot_zd, 2) if not np.isnan(d.pivot_zd) else '',
                '笔方向': stroke_map.get(d.stroke_direction, str(d.stroke_direction)),
                '线段方向': seg_map.get(d.segment_direction, str(d.segment_direction)),
                '买点类型': buy_map.get(d.buy_point, str(d.buy_point)),
                '买点置信度': round(d.buy_confidence, 3) if d.buy_point > 0 else '',
                '卖点类型': sell_map.get(d.sell_point, str(d.sell_point)),
                '卖点置信度': round(d.sell_confidence, 3) if d.sell_point > 0 else '',
                '二买确认': '是' if d.second_buy else '',
                '二买置信度': round(d.second_buy_conf, 3) if d.second_buy else '',
                '双级确认买': '是' if d.confirmed_buy else '',
                '双级确认卖': '是' if d.confirmed_sell else '',
                '信号级别': lvl_map.get(d.signal_level, str(d.signal_level)),
                '买入强度': round(d.buy_strength, 3) if d.buy_strength > 0 else '',
                '卖出强度': round(d.sell_strength, 3) if d.sell_strength > 0 else '',
                '底背驰': '是' if d.bottom_divergence else '',
                '顶背驰': '是' if d.top_divergence else '',
                '均线排列': round(d.alignment_score, 3),
                '均线破位': ', '.join(d.ma_breakdowns) if d.ma_breakdowns else '',
                '量异常': d.volume_anomaly if d.volume_anomaly != '正常' else '',
                '笔趋势耗尽': '是' if d.bi_trend_depletion else '',
                '综合评分': round(d.composite_score, 4),
                '机会等级': d.opportunity_rank,
                '风险等级': d.risk_level,
                '止损价': round(d.stop_loss_price, 2),
            })

        df = pd.DataFrame(rows)
        csv_path = os.path.join(self.output_dir, f'缠论全量分析_{date_str}.csv')
        df.to_csv(csv_path, index=False, encoding='utf-8-sig')
        print(f'  全量缠论分析表: {len(df)} 只股票')

        return csv_path

    def _export_word(self, md_path: str, date_str: str, suffix: str = '') -> str:
        """将Markdown报告转为Word文档，方便用户自行调整"""
        try:
            from docx import Document
            from docx.shared import Pt, Inches, RGBColor, Cm
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.oxml.ns import qn
        except ImportError:
            print('  [WARN] python-docx未安装, 跳过Word导出 (pip install python-docx)')
            return ''

        if not os.path.exists(md_path):
            return ''

        with open(md_path, 'r', encoding='utf-8') as f:
            md_content = f.read()

        doc = Document()
        # 设置默认字体
        style = doc.styles['Normal']
        font = style.font
        font.name = '微软雅黑'
        font.size = Pt(10)
        style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

        lines = md_content.split('\n')
        i = 0
        while i < len(lines):
            line = lines[i]

            # 空行
            if not line.strip():
                i += 1
                continue

            # 表格
            if line.strip().startswith('|') and line.strip().endswith('|'):
                table_lines = []
                while i < len(lines) and lines[i].strip().startswith('|'):
                    table_lines.append(lines[i])
                    i += 1
                self._add_word_table(doc, table_lines)
                continue

            # 标题
            if line.startswith('# '):
                p = doc.add_heading(line[2:].strip(), level=1)
            elif line.startswith('## '):
                p = doc.add_heading(line[3:].strip(), level=2)
            elif line.startswith('### '):
                p = doc.add_heading(line[4:].strip(), level=3)
            elif line.startswith('#### '):
                p = doc.add_heading(line[5:].strip(), level=4)
            # 引用
            elif line.startswith('> '):
                p = doc.add_paragraph(line[2:].strip())
                p.style = doc.styles['Normal']
                p.runs[0].font.italic = True
                p.runs[0].font.color.rgb = RGBColor(100, 100, 100)
            # 分隔线
            elif line.strip() == '---':
                doc.add_paragraph('─' * 60)
            # 代码块
            elif line.strip() == '```':
                i += 1
                code_lines = []
                while i < len(lines) and lines[i].strip() != '```':
                    code_lines.append(lines[i])
                    i += 1
                if code_lines:
                    p = doc.add_paragraph()
                    p.style = doc.styles['Normal']
                    run = p.add_run('\n'.join(code_lines))
                    run.font.name = 'Courier New'
                    run.font.size = Pt(7)
                i += 1
                continue
            # 列表
            elif line.strip().startswith('- ') or line.strip().startswith('* '):
                p = doc.add_paragraph(line.strip()[2:], style='List Bullet')
            # 普通段落（处理 **粗体**）
            else:
                p = doc.add_paragraph()
                self._add_formatted_run(p, line.strip())

            i += 1

        # 保存
        fname = f'daily_review_{date_str}{suffix}.docx'
        word_path = os.path.join(self.output_dir, fname)
        doc.save(word_path)
        print(f'  Word文档: {fname}')
        return word_path

    def _add_word_table(self, doc, table_lines):
        """解析Markdown表格行并添加到Word文档"""
        try:
            from docx.shared import Pt
        except ImportError:
            return
        rows = []
        for tl in table_lines:
            cells = [c.strip() for c in tl.strip().split('|')[1:-1]]
            # 跳过分隔行
            if all(c.startswith('---') or c.startswith('===') or set(c) <= set('-:= ')
                   for c in cells):
                continue
            rows.append(cells)

        if not rows:
            return

        ncols = max(len(r) for r in rows)
        # 补齐列数
        for r in rows:
            while len(r) < ncols:
                r.append('')

        table = doc.add_table(rows=len(rows), cols=ncols, style='Table Grid')
        table.autofit = True

        for ri, row_cells in enumerate(rows):
            for ci, cell_text in enumerate(row_cells):
                cell = table.cell(ri, ci)
                # 清空默认段落内容
                for p in cell.paragraphs:
                    p.clear()
                p = cell.paragraphs[0]
                # 判断是否表头
                is_header = (ri == 0 and any(kw in str(rows[0]) for kw in ['代码', '名称', '排名', '指标', '行业', '日期', '操作', '买点', '卖点']))
                try:
                    run = p.add_run(str(cell_text))
                    run.font.size = Pt(8)
                    run.font.name = '微软雅黑'
                    run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
                    if is_header:
                        run.bold = True
                except Exception:
                    p.text = str(cell_text)
                p.alignment = 1  # WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph()  # 表后空行

    def _add_formatted_run(self, paragraph, text):
        """解析**粗体**和![图片]()并添加到段落"""
        import re
        try:
            from docx.shared import RGBColor
        except ImportError:
            RGBColor = None
        # 先处理图片占位
        img_match = re.search(r'!\[.*?\]\(charts/(\d+)_Kline\.png\)', text)
        if img_match:
            code = img_match.group(1)
            run = paragraph.add_run(f'【📷 请在此插入 {code} K线图：charts/{code}_Kline.png】')
            run.bold = True
            if RGBColor:
                run.font.color.rgb = RGBColor(200, 50, 50)
            return

        parts = re.split(r'(\*\*.*?\*\*)', text)
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                run = paragraph.add_run(part[2:-2])
                run.bold = True
            else:
                paragraph.add_run(part)

    def _export_document(self, report: DailyReviewReport, path: str):
        """导出完整复盘报告文档 (Markdown)"""
        m = report.market
        summary = report.summary
        lines = []

        def w(text=''):
            lines.append(text)

        # 动态标题：根据市场状态生成
        sentiment_keywords = {
            '冰点': '冰点情绪', '修复': '情绪修复', '高潮': '情绪高潮',
            '分歧': '市场分歧', '退潮': '情绪退潮',
        }
        regime_keywords = {'牛市': '牛市格局', '震荡': '震荡格局', '熊市': '熊市防御'}
        sk = sentiment_keywords.get(m.sentiment_cycle, '')
        rk = regime_keywords.get(m.regime, '')
        # 关键主题
        b3_count = sum(1 for d in report.all_diagnoses if d.buy_point == 3)
        main_lines = [ind for ind in report.main_lines if ind.summary.startswith('★')]
        top_ind = main_lines[0].industry if main_lines else ''
        if b3_count > 30:
            theme = f'聚焦{top_ind}等主线三买机会' if top_ind else '三买机会密集'
        elif m.sentiment_cycle == '冰点':
            theme = '等待冰点后的结构信号'
        elif m.sentiment_cycle == '高潮':
            theme = '高潮中关注分歧信号'
        else:
            theme = f'结构分化中的机会挖掘'

        title = f'{report.date} 每日复盘：{sk}下的{rk}，{theme}'
        w(f'# {title}')
        w()

        # ─── 开场白 ───
        # 根据市场状态动态生成开场判断
        sentiment_analysis = {
            '冰点': '情绪已降至冰点，上涨家数稀少。但冰点往往也是调整末端——牛市中的冰点意味着反弹动能在积蓄，密切关注率先企稳的方向。',
            '修复': '情绪正在修复，上涨家数开始增加，但尚未形成一致预期。当前是识别主线的关键窗口，重点关注率先放量领涨的板块。',
            '高潮': '情绪高涨，涨停众多，赚钱效应强。但高潮之后往往伴随分歧，追高需谨慎，应关注板块轮动节奏和个股买卖点结构。',
            '分歧': '市场进入分歧阶段，涨跌互现，多空胶着。此时结构清晰、有买点支撑的个股更抗跌，应回避方向不明的标的。',
            '退潮': '市场进入退潮期，涨停减少，亏钱效应显现。建议控制仓位、谨慎开新仓，等待情绪冰点后的修复信号。',
        }.get(m.sentiment_cycle, '市场方向尚不明确，需等待更清晰的结构信号择向。')

        regime_note = {
            '牛市': '指数处于牛市结构，中期趋势向上，回调是布局机会。',
            '震荡': '指数处于震荡格局，方向有待选择，以波段操作为主。',
            '熊市': '⚠ 指数处于熊市结构，需严格控制仓位，反弹减仓为主。',
        }.get(m.regime, '')

        vol_analysis = {
            '放量': '今日放量，多空交战激烈。',
            '缩量': '今日缩量，市场参与度下降。',
            '平量': '今日量能持平，市场维持正常换手。',
        }.get(m.index_volume, '')

        risk_warning = ' ⚠ 系统检测到熊市风险，需严格控制仓位。' if m.bear_risk else ' 系统未检测到极端风险信号。'

        # 买卖对比判断
        buy_c = summary["buy_signals"]
        sell_c = summary["sell_signals"]
        if buy_c > sell_c * 1.5:
            bias = '买点信号显著多于卖点，做多动能占优，可积极选股。'
        elif sell_c > buy_c * 1.5:
            bias = '卖点信号显著多于买点，市场偏空，以防御为主。'
        elif buy_c > sell_c:
            bias = '买点略多于卖点，结构性机会存在，精选个股参与。'
        else:
            bias = '买卖信号接近，市场处于均衡状态，等待方向选择。'

        w(f'今天是{report.date}，我们对全市场{summary["total_stocks"]}只个股进行系统性复盘。')
        w()
        w(f'**大盘定调**: 上证指数收于{m.index_close:.1f}点（{m.index_change_pct:+.2f}%）。'
          f'{regime_note}')
        w(f'**情绪与量能**: {m.sentiment_cycle}周期。{vol_analysis}{risk_warning}')
        w(f'**买卖格局**: 全市场买点信号{buy_c}只，卖点信号{sell_c}只。{bias}')
        w(f'**结构特征**: 上涨趋势{summary["up_trend"]}只（{summary["up_trend"]/max(summary["total_stocks"],1)*100:.1f}%），'
          f'盘整{summary["consolidation"]}只，下跌趋势{summary["down_trend"]}只。'
          f'涨停{report.limit_up_summary.total_limit_up if report.limit_up_summary else 0}只，'
          f'跌停{m.limit_down_count}只。')
        w()
        w('---')
        w()

        # ─── 1. 大盘体检 ───
        w('## 一、大盘环境深度体检')
        w()
        w(f'| 指标 | 数值 | 指标 | 数值 |')
        w(f'|------|------|------|------|')
        w(f'| 上证指数 | {m.index_close:.1f} | 涨跌幅 | {m.index_change_pct:+.2f}% |')
        w(f'| 量能 | {m.index_volume} ({m.volume_ratio:.1f}x均量) | 情绪周期 | {m.sentiment_cycle} |')
        w(f'| 市场状态 | {m.regime} (置信度 {m.regime_confidence:.0%}) | 风格 | {m.style_regime} |')
        w(f'| 动量 | {m.momentum_score:+.2f} | 波动率 | {m.volatility_pct:.1%} |')
        w(f'| 上涨/下跌 | {m.advancing} / {m.declining} ({m.up_pct:.0f}%) | 涨停/跌停 | {m.limit_up_count} / {m.limit_down_count} |')
        w(f'| 综合判定 | {m.summary} | 熊市风险 | {"⚠ 是" if m.bear_risk else "✓ 否"} |')
        w()

        # ─── 2. 行业主线 ───
        w('## 二、行业主线分析')
        w()
        if report.main_lines:
            w(f'| 排名 | 行业 | 股票数 | 上涨趋势 | 下跌趋势 | B1/B2/B3 | 梯队 | 综合评分 | 判定 |')
            w(f'|------|------|--------|----------|----------|----------|------|----------|------|')
            for ind in report.main_lines[:20]:
                formation_bar = '★' * int(ind.formation_score * 3) + '☆' * (3 - int(ind.formation_score * 3))
                w(f'| {ind.hot_rank} | {ind.industry} | {ind.stock_count} | '
                  f'{ind.up_trend_pct:.0f}% | {ind.down_trend_pct:.0f}% | '
                  f'{ind.b1_count}/{ind.b2_count}/{ind.b3_count} | {formation_bar} | '
                  f'{ind.avg_composite_score:+.2f} | {ind.summary} |')
            w()

            # 主线行业个股明细（TOP5行业列出具体股票）
            top_inds = [ind for ind in report.main_lines
                       if ind.summary.startswith('★')][:5]
            if top_inds:
                w('### 主线行业个股明细')
                w()

                # 构建 industry → [diagnoses] 映射
                ind_diag = defaultdict(list)
                for d in report.all_diagnoses:
                    if d.industry:
                        ind_diag[d.industry].append(d)

                trend_map = {2: '上涨趋势', 1: '盘整', 0: '-', -2: '下跌趋势'}
                pp_map = {1: '中枢上', 0: '中枢内', -1: '中枢下'}
                bp_label = {1: 'B1', 2: 'B2', 3: 'B3'}

                for ind in top_inds:
                    stocks = ind_diag.get(ind.industry, [])
                    # 筛选有买点的股票
                    buy_stocks = [s for s in stocks if s.buy_point > 0]
                    buy_stocks.sort(key=lambda x: x.composite_score, reverse=True)
                    if not buy_stocks:
                        continue

                    w(f'**{ind.industry}**（梯队评分 {ind.formation_score:.1f}/3.0）— '
                      f'{ind.summary.lstrip("★○✗— ")}。'
                      f'板块内上涨率{ind.up_trend_pct:.0f}%，'
                      f'拥有B3龙头{ind.b3_count}只、B1后备{ind.b1_count}只。')
                    w()
                    w(f'| 代码 | 名称 | 买点 | 收盘 | 涨跌 | 走势 | 中枢位 | 综合 |')
                    w(f'|------|------|------|------|------|------|--------|------|')
                    for s in buy_stocks[:8]:
                        bp = bp_label.get(s.buy_point, '')
                        if s.second_buy:
                            bp = 'B2*'
                        w(f'| {s.symbol} | {s.name} | {bp} | {s.close:.2f} | '
                          f'{s.change_pct:+.1f}% | {trend_map.get(s.trend_type, "?")} | '
                          f'{pp_map.get(s.pivot_position, "?")} | {s.composite_score:+.2f} |')
                    w()
        else:
            w('无数据')
        w()

        # ─── 3. 涨停板缠论归因 ───
        if report.limit_up_summary and report.limit_up_summary.limit_up_stocks:
            lu = report.limit_up_summary
            w('## 三、涨停板缠论归因分析')
            w()
            w(f'- **涨停总数**: {lu.total_limit_up} 只')
            w(f'- **主导模式**: {lu.dominant_pattern}')
            w(f'- **类型分布**: {", ".join(f"{t}:{c}只" for t, c in sorted(lu.type_distribution.items(), key=lambda x: x[1], reverse=True))}')

            # 个股归因表
            w()
            w(f'| 代码 | 名称 | 涨幅 | 量比 | 归因类型 | 置信度 | 涨停前结构 |')
            w(f'|------|------|------|------|----------|--------|------------|')
            trend_map = {2: '上涨趋势', 1: '盘整', 0: '--', -2: '下跌趋势'}
            pp_map = {1: '中枢上', 0: '中枢内', -1: '中枢下'}
            for d in lu.limit_up_stocks:
                pre = f'{trend_map.get(d.pre_trend_type, "?")}/{pp_map.get(d.pre_pivot_position, "?")}'
                w(f'| {d.symbol} | {d.name} | {d.change_pct:+.1f}% | {d.vol_spike:.1f}x | '
                  f'{d.limit_up_type} | {d.type_confidence:.0%} | {pre} |')

            # 缠论归因详情
            w()
            w('### 缠论归因详情')
            for d in lu.limit_up_stocks:
                w(f'- **{d.symbol} {d.name}** [{d.limit_up_type}] (置信度 {d.type_confidence:.0%})')
                w(f'  - {d.chan_explanation}')
            w()

            # 共性总结
            if lu.common_traits:
                w('### 涨停板共性总结')
                for i, trait in enumerate(lu.common_traits):
                    w(f'{i+1}. {trait}')
                w()

            # 行业集中
            if lu.industry_concentration:
                w('### 涨停行业集中度')
                w(f'| 行业 | 涨停数 |')
                w(f'|------|--------|')
                for ind, cnt in lu.industry_concentration.items():
                    w(f'| {ind} | {cnt} |')
                w()

        # ─── 4. 买卖点机会 ───
        w('## 四、缠论买卖点机会')
        w()

        # 买点类型概览
        buy_descriptions = {
            'A_strong_buy': '笔线段双级别强共振，信号可靠性最高',
            'B2_confirmed': '一买后回踩不破前低，最安全的加仓点(第49课)',
            'B1_reversal': '下跌趋势末端底背驰，左侧抄底(第21课)',
            'B3_acceleration': '突破中枢上沿回踩确认，主升浪起点(第54课)',
            'double_confirmed': '双周期信号确认，可靠性高于单级别',
        }
        op_labels = {
            'A_strong_buy': 'A级 强买（多级别确认）',
            'B2_confirmed': 'B2 二买确认',
            'B1_reversal': 'B1 一买（底部反转）',
            'B3_acceleration': 'B3 三买（趋势加速）',
            'double_confirmed': '双级别确认',
        }
        w(f'| 买点类型 | 数量 | 说明 |')
        w(f'|----------|------|------|')
        for key, label in op_labels.items():
            snaps = report.opportunities.get(key, [])
            if snaps:
                w(f'| {label} | {len(snaps)}只 | {buy_descriptions.get(key, "")} |')
        w()

        # 个股明细
        trend_map = {2: '上涨趋势', 1: '盘整', 0: '--', -2: '下跌趋势'}
        pp_map = {1: '上方', 0: '内部', -1: '下方'}
        bp_map = {1: 'B1', 2: 'B2', 3: 'B3'}
        SHOW_DETAIL_MAX = 10

        for key, label in op_labels.items():
            snaps = report.opportunities.get(key, [])
            if not snaps:
                continue

            n_total = len(snaps)
            if n_total <= SHOW_DETAIL_MAX:
                w(f'### {label}（共{n_total}只，全部如下）')
                w()
                w(f'| 代码 | 名称 | 行业 | 收盘 | 涨跌 | 买点 | 走势 | 中枢位 | 置信 | 综合 |')
                w(f'|------|------|------|------|------|------|------|--------|------|------|')
                for s in snaps:
                    bp = bp_map.get(s.buy_point, '')
                    if s.second_buy: bp = 'B2*'
                    w(f'| {s.symbol} | {s.name} | {s.industry} | {s.close:.2f} | '
                      f'{s.change_pct:+.1f}% | {bp} | {trend_map.get(s.trend_type, "?")} | '
                      f'{pp_map.get(s.pivot_position, "?")} | {s.buy_confidence:.0%} | '
                      f'{s.composite_score:+.2f} |')
                w()
            else:
                w(f'### {label}（共{n_total}只，精选前{SHOW_DETAIL_MAX}只，全部代码见下方）')
                w()
                w(f'| 代码 | 名称 | 行业 | 收盘 | 涨跌 | 买点 | 走势 | 中枢位 | 置信 | 综合 |')
                w(f'|------|------|------|------|------|------|------|--------|------|------|')
                for s in snaps[:SHOW_DETAIL_MAX]:
                    bp = bp_map.get(s.buy_point, '')
                    if s.second_buy: bp = 'B2*'
                    w(f'| {s.symbol} | {s.name} | {s.industry} | {s.close:.2f} | '
                      f'{s.change_pct:+.1f}% | {bp} | {trend_map.get(s.trend_type, "?")} | '
                      f'{pp_map.get(s.pivot_position, "?")} | {s.buy_confidence:.0%} | '
                      f'{s.composite_score:+.2f} |')
                w()
                all_codes = [s.symbol for s in snaps]
                w(f'**全部{n_total}只代码**: {", ".join(all_codes)}')
                w()

        # ─── 5. 风险预警 ───
        w('## 五、风险预警')
        w()

        # 风险类型概览
        risk_labels = {
            'confirmed_sell': '多级别确认卖点',
            'S1_top_reversal': 'S1 一卖（顶部反转）',
            'S3_breakdown': 'S3 三卖（趋势破位）',
            'top_divergence': '顶背驰预警',
            'ma_breakdown': '均线破位（≥2条）',
        }
        risk_descriptions = {
            'confirmed_sell': '笔线段双级别共振卖点，下跌概率最高',
            'S1_top_reversal': '上涨趋势末端，顶背驰确认，多头力竭',
            'S3_breakdown': '跌破中枢下沿，反抽不涨回，趋势破位',
            'top_divergence': 'MACD顶背离预警，可能尚未触发卖点',
            'ma_breakdown': '股价跌破关键均线支撑（MA5/10/20/60中≥2条）',
        }
        w(f'| 风险类型 | 数量 | 说明 |')
        w(f'|----------|------|------|')
        for key, label in risk_labels.items():
            snaps = report.risk_alerts.get(key, [])
            if snaps:
                w(f'| {label} | {len(snaps)}只 | {risk_descriptions.get(key, "")} |')
        w()

        # 个股明细
        pp_map = {1: '上方', 0: '内部', -1: '下方'}
        sp_map = {1: 'S1', 2: 'S2', 3: 'S3'}
        SHOW_DETAIL_MAX = 10  # 超过此数量时，前N只详细展示，其余仅列代码

        for key, label in risk_labels.items():
            snaps = report.risk_alerts.get(key, [])
            if not snaps:
                continue

            n_total = len(snaps)
            if n_total <= SHOW_DETAIL_MAX:
                # 全部详细展示
                w(f'### {label}（共{n_total}只，全部如下）')
                w()
                w(f'| 代码 | 名称 | 行业 | 收盘 | 涨跌 | 卖点 | 走势 | 中枢位 | 置信 | 综合 |')
                w(f'|------|------|------|------|------|------|------|--------|------|------|')
                for s in snaps:
                    conf = s.sell_confidence
                    sp = sp_map.get(s.sell_point, '')
                    w(f'| {s.symbol} | {s.name} | {s.industry} | {s.close:.2f} | '
                      f'{s.change_pct:+.1f}% | {sp} | {trend_map.get(s.trend_type, "?")} | '
                      f'{pp_map.get(s.pivot_position, "?")} | {conf:.0%} | '
                      f'{s.composite_score:+.2f} |')
                w()
            else:
                # 精选TOP10 + 全部代码列表
                w(f'### {label}（共{n_total}只，精选前{SHOW_DETAIL_MAX}只，全部代码见下方）')
                w()
                w(f'| 代码 | 名称 | 行业 | 收盘 | 涨跌 | 卖点 | 走势 | 中枢位 | 置信 | 综合 |')
                w(f'|------|------|------|------|------|------|------|--------|------|------|')
                for s in snaps[:SHOW_DETAIL_MAX]:
                    conf = s.sell_confidence
                    sp = sp_map.get(s.sell_point, '')
                    w(f'| {s.symbol} | {s.name} | {s.industry} | {s.close:.2f} | '
                      f'{s.change_pct:+.1f}% | {sp} | {trend_map.get(s.trend_type, "?")} | '
                      f'{pp_map.get(s.pivot_position, "?")} | {conf:.0%} | '
                      f'{s.composite_score:+.2f} |')
                w()
                # 全部代码压缩列表（8列一行）
                all_codes = [s.symbol for s in snaps]
                w(f'**全部{n_total}只代码**: {", ".join(all_codes)}')
                w()

        # ─── 图表 ─── (自动生成 + 截图占位)
        w('## 六、重点个股截图占位（请手动截图替换）')
        w()
        w('> 请在交易软件中截取以下个股的日K线图（含MA5/MA10/MA20/MA60均线 + MACD指标 + 成交量），')
        w('> 保存为 `charts/XXXXXX_Kline.png` 后替换下方占位符。')
        w()

        # 收集需要截图的重点标的
        chart_stocks = []
        seen = set()
        for d in report.opportunities.get('A_strong_buy', [])[:3]:
            if d.symbol not in seen:
                chart_stocks.append((d, 'A级买点'))
                seen.add(d.symbol)
        for d in report.opportunities.get('B2_confirmed', [])[:2]:
            if d.symbol not in seen:
                chart_stocks.append((d, 'B2二买确认'))
                seen.add(d.symbol)
        for d in report.opportunities.get('B1_reversal', [])[:3]:
            if d.symbol not in seen:
                chart_stocks.append((d, 'B1一买'))
                seen.add(d.symbol)
        for d in report.opportunities.get('B3_acceleration', [])[:2]:
            if d.symbol not in seen:
                chart_stocks.append((d, 'B3三买'))
                seen.add(d.symbol)

        if chart_stocks:
            w(f'| # | 代码 | 名称 | 行业 | 收盘 | 涨跌 | 类型 | 综合 | 截图位置 |')
            w(f'|---|------|------|------|------|------|------|------|----------|')
            for i, (d, reason) in enumerate(chart_stocks):
                img_placeholder = f'![{d.symbol}](charts/{d.symbol}_Kline.png)'
                w(f'| {i+1} | {d.symbol} | {d.name} | {d.industry} | {d.close:.2f} | '
                  f'{d.change_pct:+.1f}% | {reason} | {d.composite_score:+.3f} | '
                  f'{img_placeholder} |')
            w()
            w('**截图要求**: 日K线图，含MA5/MA10/MA20/MA60均线，MACD指标（含DIF/DEA/柱），成交量。标注买卖点位置和中枢区间。')
        else:
            w('暂无需要截图的重点标的。')
        w()

        # ─── 7. 持仓评估与明日调仓计划 ───
        w('## 七、持仓评估与明日调仓计划')
        w()

        # 仓位建议
        if m.sentiment_cycle in ('冰点', '退潮'):
            pos_advice = '防御仓位 (1-3成)'
        elif m.sentiment_cycle == '高潮':
            pos_advice = '积极仓位 (5-7成)'
        elif m.bear_risk:
            pos_advice = '轻仓或空仓 (0-2成)'
        else:
            pos_advice = '中性仓位 (3-5成)'
        w(f'**仓位建议**: {pos_advice} | **可用资金**: ¥{report.portfolio_cash:,.0f}')
        w()

        # --- 当前持仓评估 ---
        holdings = report.portfolio_holdings
        if holdings:
            total_mv = sum(h['market_value'] for h in holdings)
            total_pnl = total_mv - sum(h['cost']*h['shares'] for h in holdings)
            w('### 当前持仓评估')
            w()
            w(f'持仓 {len(holdings)} 只，总市值 ¥{total_mv:,.0f}，浮动盈亏 ¥{total_pnl:+,.0f}')
            w()
            w(f'| 代码 | 名称 | 持仓(股) | 成本 | 现价 | 市值 | 盈亏% | 走势 | 买点 | 卖点 | 风险 | 操作建议 |')
            w(f'|------|------|----------|------|------|------|-------|------|------|------|------|----------|')
            for h in holdings:
                w(f'| {h["code"]} | {h["name"]} | {h["shares"]:,.0f} | '
                  f'{h["cost"]:.2f} | {h["close"]:.2f} | {h["market_value"]:,.0f} | '
                  f'{h["pnl_pct"]:+.1f}% | {h["trend"]} | {h["buy_pt"]} | {h["sell_pt"]} | '
                  f'{h["risk"]} | {h["action"]} |')
            w()
            # 逐只深度分析
            for h in holdings:
                w(f'#### {h["code"]} {h["name"]}')
                w()
                w(f'| 项目 | 内容 |')
                w(f'|------|------|')
                w(f'| 收盘价 | {h["close"]:.2f} ({h["change_pct"]:+.1f}%) |')
                w(f'| 持仓 | {h["shares"]}股, 成本{h["cost"]:.2f}, '
                  f'市值¥{h["market_value"]:,.0f}, 盈亏{h["pnl_pct"]:+.1f}% |')
                w(f'| 缠论结构 | {h["structure_desc"]} |')
                w(f'| 买卖点 | 买点={h["buy_pt"]}, 卖点={h["sell_pt"]} |')
                w(f'| 背驰 | {h["div_desc"]} |')
                w(f'| 量价 | {h["vol_desc"]} |')
                w(f'| 综合评分 | {h["composite"]:+.2f} ({h["rank"]}级), 风险={h["risk"]} |')
                w(f'| 关键价位 | {h["key_levels"]} |')
                w(f'| 风险信号 | {h["risk_detail"]} |')
                if h["take_profit"]:
                    w(f'| 止盈参考 | {h["take_profit"]} |')
                w(f'| 操作建议 | **{h["action"]}** — {h["reason"]} |')
                w()
        else:
            w('### 当前持仓: 空仓')
            w()

        # --- 持仓替换计划 ---
        replacements = report.replacement_plan
        if replacements:
            w('### 调仓替换计划')
            w()
            w(f'| 持仓股 | 操作 | 原因 | 替换标的 | 评分 | 现价 | 紧急度 |')
            w(f'|--------|------|------|----------|------|------|--------|')
            for r in replacements:
                w(f'| {r["code"]} {r["name"]} | **{r["action"]}** | {r["reason"]} | '
                  f'{r["replacement_code"]} {r["replacement_name"]} | '
                  f'{r["replacement_score"]:.3f} | {r["replacement_price"]:.2f} | '
                  f'{r["urgency"]} |')
            w()

        # --- 当日选股结果 ---
        selections = report.today_selections
        if selections:
            w('### 今日系统选股结果')
            w()
            w(f'系统选出 {len(selections)} 只标的:')
            w()
            w(f'| # | 代码 | 名称 | 行业 | 评分 | 权重 | 目标仓位 | 现价 | 信号 |')
            w(f'|---|------|------|------|------|------|----------|------|------|')
            for i, s in enumerate(selections):
                code = str(s.get('股票代码', ''))
                name = str(s.get('股票名称', ''))
                ind = str(s.get('行业', ''))
                score = s.get('综合评分', 0)
                weight = s.get('权重', 0)
                target = s.get('目标仓位', 0)
                price = s.get('当前价格', 0)
                signal = str(s.get('信号强度', ''))
                w(f'| {i+1} | {code} | {name} | {ind} | '
                  f'{float(score):.3f} | {float(weight)*100:.1f}% | '
                  f'¥{float(target):,.0f} | {float(price):.2f} | {signal} |')
            w()

            # 调仓对比
            if holdings:
                hold_codes = {h['code'] for h in holdings}
                sel_codes = {str(s.get('股票代码', '')) for s in selections}
                to_buy = sel_codes - hold_codes
                to_sell = hold_codes - sel_codes
                to_hold = sel_codes & hold_codes

                w('### 调仓建议')
                w()
                if to_buy:
                    w(f'**新买入** ({len(to_buy)}只): {", ".join(sorted(to_buy))}')
                if to_sell:
                    w(f'**清仓卖出** ({len(to_sell)}只): {", ".join(sorted(to_sell))}')
                if to_hold:
                    w(f'**继续持有** ({len(to_hold)}只): {", ".join(sorted(to_hold))}')
                if not to_buy and not to_sell:
                    w('无需调仓，维持现有持仓。')
                w()
        else:
            w('### 今日选股: 未运行')
            w()
            w('> 提示: 运行 `python export_selection.py` 生成选股结果后，报告中会自动包含调仓计划。')
            w()

        # --- 作战计划（策略选股+缠论诊断交叉分析） ---
        w('### 作战计划（策略选股 × 缠论诊断）')
        w()
        if report.portfolio_suggestions:
            w(f'| # | 代码 | 名称 | 行业 | 收盘 | 买点 | 止损 | 综合 | 等级 |')
            w(f'|---|------|------|------|------|------|------|------|------|')
            for i, d in enumerate(report.portfolio_suggestions):
                bp = {1: 'B1', 2: 'B2', 3: 'B3'}.get(d.buy_point, '')
                if d.second_buy:
                    bp = 'B2*'
                w(f'| {i+1} | {d.symbol} | {d.name} | {d.industry} | {d.close:.2f} | '
                  f'{bp} | {d.stop_loss_price:.2f} | {d.composite_score:+.2f} | {d.opportunity_rank} |')
        else:
            w('无符合条件的推荐标的')
        w()

        # ─── 8. 系统进化 ───
        pa = report.pred_accuracy
        ah = report.accuracy_history

        # 准确率历史趋势
        has_history = len(ah) >= 1
        has_current = pa and pa.get('available')

        if not has_history and not has_current:
            w('## 八、系统准确率进化记录')
            w()
            w('> 准确率追踪尚未开始。')
            w('> 使用 `python analysis/chan_review.py --evolve` 启动自动追踪。')
            w('> 系统会在每次复盘后自动评估前一日推荐标的的实际表现。')
            w()
        elif len(ah) >= 2:
            w('## 八、系统准确率进化记录')
            w()
            w('> 每次复盘自动评估前一日推荐标的的实际表现，记录方向准确率和超额收益。')
            w('> 观察准确率趋势可判断系统参数调整是否有效。')
            w()
            w(f'| 日期 | 回测日 | 方向准确率 | 平均收益 | 超额收益 | 信号数 |')
            w(f'|------|--------|------------|----------|----------|--------|')
            for h in ah[-20:]:  # 最近20条
                acc_icon = '✓' if h['direction_accuracy'] >= 0.5 else '✗'
                ex_icon = '+' if h['excess_return'] > 0 else ''
                w(f'| {h["date"]} | {h.get("prev_date", "?")} | '
                  f'{acc_icon} {h["direction_accuracy"]:.0%} | '
                  f'{h["avg_return"]:+.2f}% | '
                  f'{ex_icon}{h["excess_return"]:+.2f}% | '
                  f'{h["hits"]}/{h["total_signals"]} |')

            # 趋势统计
            recent_accs = [h['direction_accuracy'] for h in ah[-5:]]
            if len(recent_accs) >= 3:
                avg_recent = sum(recent_accs) / len(recent_accs)
                first_half = [h['direction_accuracy'] for h in ah[:len(ah)//2]]
                last_half = [h['direction_accuracy'] for h in ah[len(ah)//2:]]
                avg_first = sum(first_half) / len(first_half) if first_half else 0
                avg_last = sum(last_half) / len(last_half) if last_half else 0
                trend = '↑ 上升' if avg_last > avg_first + 0.05 else ('↓ 下降' if avg_first > avg_last + 0.05 else '→ 持平')
                w()
                w(f'**准确率趋势**: 近5日均{avg_recent:.0%} | '
                  f'前半段均{avg_first:.0%} → 后半段均{avg_last:.0%} {trend}')
            w()

        elif len(ah) == 1:
            h = ah[0]
            w('## 八、系统准确率进化记录')
            w()
            w(f'| 日期 | 回测日 | 方向准确率 | 平均收益 | 超额收益 | 信号数 |')
            w(f'|------|--------|------------|----------|----------|--------|')
            acc_icon = '✓' if h['direction_accuracy'] >= 0.5 else '✗'
            ex_icon = '+' if h['excess_return'] > 0 else ''
            w(f'| {h["date"]} | {h.get("prev_date", "?")} | '
              f'{acc_icon} {h["direction_accuracy"]:.0%} | '
              f'{h["avg_return"]:+.2f}% | '
              f'{ex_icon}{h["excess_return"]:+.2f}% | '
              f'{h["hits"]}/{h["total_signals"]} |')
            w()
            w('> 历史数据积累中，使用 `--evolve` 参数每次复盘后自动追加。')
            w()

        elif pa and pa.get('available'):
            w('## 八、预测准确性回测')
            w()
            w(f'- **回测日期**: {pa.get("prev_date", "?")}')
            w(f'- **方向准确率**: {pa["direction_accuracy"]:.0%} ({pa["hits"]}/{pa["total_buy_signals"]})')
            w(f'- **平均收益**: {pa["avg_return"]:+.2f}%')
            w(f'- **市场基准**: {pa.get("market_return", 0):+.2f}%')
            w(f'- **超额收益**: {pa.get("excess_return", 0):+.2f}%')
            by_type = pa.get('by_type', {})
            if by_type:
                w('- **按买点类型**:')
                for bt, info in sorted(by_type.items()):
                    w(f'  - {bt}: {info["accuracy"]:.0%} ({info["count"]}次)')
            w()

        if report.evolution_patches:
            w('## 九、系统自进化建议')
            w()
            for p in report.evolution_patches:
                urgency_emoji = {'urgent': '🔴', 'high': '🟠', 'medium': '🟡', 'low': '🟢'}
                w(f'- {urgency_emoji.get(p.urgency, "⚪")} **[{p.section}.{p.key}]**')
                w(f'  - 当前值: `{p.current_value}` → 建议值: `{p.suggested_value}`')
                w(f'  - 理由: {p.reason}')
                w(f'  - 置信度: {p.confidence:.0%}')
            w()

        # ─── 10. 复盘反思 ───
        w('## 十、复盘反思')
        w()
        # 当日关键数据盘点
        w('### 今日市场要点')
        w()
        w(f'- 市场状态: {m.regime}，情绪周期: {m.sentiment_cycle}，风格偏向: {m.style_regime}')
        w(f'- 全市场买点信号 {summary["buy_signals"]} 只（B1一买: {sum(1 for d in report.all_diagnoses if d.buy_point==1)}, '
          f'B2二买: {summary["second_buy"]}, B3三买: {sum(1 for d in report.all_diagnoses if d.buy_point==3)}）')
        w(f'- 卖点信号 {summary["sell_signals"]} 只，高危股票 {summary["high_risk"]} 只')
        w(f'- A级机会 {summary["a_rank"]} 只，B级机会 {summary["b_rank"]} 只')
        w(f'- 涨停 {report.limit_up_summary.total_limit_up if report.limit_up_summary else 0} 只，'
          f'跌停 {m.limit_down_count} 只')
        w()

        # 结构判断反思
        w('### 结构判断反思')
        w()
        # 买点浓度判断
        total = summary['total_stocks']
        buy_pct = summary['buy_signals'] / max(total, 1) * 100
        if buy_pct < 3:
            w(f'- 买点信号占比仅{buy_pct:.1f}%，市场整体做多动能不足。')
            w(f'  - 反思: 是否选股条件过严？是否需要关注结构性机会而非系统性机会？')
        elif buy_pct < 8:
            w(f'- 买点信号占比{buy_pct:.1f}%，处于正常范围，结构性机会存在。')
        else:
            w(f'- 买点信号占比{buy_pct:.1f}%，市场做多热情较高，但需警惕过度乐观。')

        # 卖点 vs 买点
        sell_count = summary['sell_signals']
        if sell_count > summary['buy_signals'] * 1.5:
            w(f'- ⚠ 卖点信号({sell_count})远超买点({summary["buy_signals"]})，市场偏空，需谨慎。')
        elif sell_count > summary['buy_signals']:
            w(f'- 卖点信号({sell_count})略多于买点({summary["buy_signals"]})，市场分歧加大。')
        elif summary['buy_signals'] > sell_count * 1.5:
            w(f'- 买点信号({summary["buy_signals"]})远超卖点({sell_count})，做多信号占优。')
        else:
            w(f'- 买卖信号接近（买{summary["buy_signals"]}/卖{sell_count}），市场处于均衡状态。')

        # 趋势分布
        up_pct = summary['up_trend'] / max(total, 1) * 100
        down_pct = summary['down_trend'] / max(total, 1) * 100
        w(f'- 趋势分布: 上涨{up_pct:.1f}% / 盘整{summary["consolidation"]/max(total,1)*100:.1f}% / 下跌{down_pct:.1f}%')
        if down_pct > up_pct * 1.5:
            w(f'  - 反思: 下跌趋势股偏多，市场整体偏弱，选股应更注重底部结构确认。')

        # 主线判断
        main_lines = [ind for ind in report.main_lines if ind.summary.startswith('★')]
        w(f'- 强主线 {len(main_lines)} 个: {", ".join(ind.industry for ind in main_lines[:5])}')
        if len(main_lines) >= 5:
            w(f'  - 反思: 主线丰富，市场有明确方向，可以积极参与。')
        elif len(main_lines) >= 2:
            w(f'  - 反思: 主线适度，在主线内择优操作，回避非主线标的。')
        else:
            w(f'  - 反思: 主线匮乏，市场缺乏共识方向，以防御为主。')

        # 系统改进方向
        w()
        w('### 系统改进备忘')
        w()
        a_count = summary['a_rank']
        b_count = summary['b_rank']
        if a_count == 0:
            w(f'- A级机会为0（共{a_count + b_count}只B级以上），信号质量偏弱。')
            w(f'  - 建议: 检查Chan理论参数是否需要调整，或当前市场不适合操作。')
        if summary['confirmed_buy'] < summary['buy_signals'] * 0.3:
            w(f'- 双级别确认买点仅{summary["confirmed_buy"]}只（占买点的{summary["confirmed_buy"]/max(summary["buy_signals"],1)*100:.0f}%），')
            w(f'  - 建议: 单级别信号过多，可能需要提高笔/线段检测的阈值。')
        w()

        # ─── 11. 缠论图解附录 ───
        w('## 十一、附：缠论三类买卖点完全图解')
        w()
        w('![三类买卖点](charts/3买3卖.png)')
        w()
        w('> 上图完整展示了缠论六种买卖点的形态结构。以下为理论要点：')
        w()

        w('### B1 第一类买点（一买）— 底背驰终结下跌趋势')
        w('- 第21课: 下跌趋势末端出现底背驰，空头力竭')
        w('- 第24课: MACD柱面积明显缩小，黄白线未同步新低')
        w('- 判定: 价格创近期新低，但对应段MACD柱面积比前一段明显缩小')
        w('- 操作: MACD金叉确认后买入，止损设在前低下方2-3%')
        w()

        w('### B2 第二类买点（二买）— 回踩不破一买低点')
        w('- 第21/49课: 一买后次级别回踩，不创新低，最安全的买点')
        w('- 判定: B1后回踩，最低点高于B1低点，成交量萎缩，MACD未破0轴')
        w('- 操作: 回踩止跌+MACD再次金叉时加仓，止损在B1低点下方')
        w()

        w('### B3 第三类买点（三买）— 突破中枢后回踩确认')
        w('- 第21/54课: 突破中枢上沿，次级别回踩不跌回中枢')
        w('- 判定: 价格突破中枢上沿，次级别回踩不跌回中枢区间，成交量放大配合')
        w('- 重要前提: 至少两个同向向上不重叠的中枢（形成趋势），底部首个中枢的B3最优')
        w('- 操作: 放量突破中枢上沿+回踩确认后介入，止损在中枢上沿下方')
        w()

        w('### S1 第一类卖点（一卖）— 顶背驰终结上涨趋势')
        w('- 第21课: 上涨趋势末端出现顶背驰，多头力竭')
        w('- 判定: 价格创近期新高，但对应段MACD柱面积比前一段明显缩小')
        w('- 操作: 顶背驰+MACD死叉时减仓或清仓')
        w()

        w('### S2 第二类卖点（二卖）— 反弹不破一卖高点')
        w('- 第21课: 一卖后反弹，不创新高，反弹无力')
        w('- 判定: S1后出现反弹，但反弹高点低于S1高点，MACD无力回到0轴上方')
        w('- 操作: 反弹无力+MACD再次死叉时清仓')
        w()

        w('### S3 第三类卖点（三卖）— 跌破中枢后反抽不涨回')
        w('- 第21课: 跌破中枢下沿，次级别反抽不涨回中枢')
        w('- 判定: 价格跌破中枢下沿，次级别反抽未能涨回中枢区间，MACD在0轴下方')
        w('- 操作: 跌破中枢下沿+反抽无力时清仓离场，不要抄底')
        w()

        # 速查表
        w('### 三类买卖点速查表')
        w()
        w(f'| 买卖点 | 定义 | 关键特征 | MACD特征 | 操作 |')
        w(f'|--------|------|----------|----------|------|')
        w(f'| B1一买 | 下跌趋势底背驰终点 | 价格新低, 背驰确认 | MACD柱面积缩小, 黄白线未同步新低 | 金叉买入, 前低止损 |')
        w(f'| B2二买 | 回踩不破一买低点 | 一买后次级别回踩 | 回踩段MACD回0轴不破 | 再次金叉加仓, B1低点止损 |')
        w(f'| B3三买 | 突破中枢上沿回踩 | 中枢上沿+次级别回踩确认 | 突破时MACD站上0轴 | 放量突破回踩买入 |')
        w(f'| S1一卖 | 上涨趋势顶背驰终点 | 价格新高, 背驰确认 | MACD柱面积缩小, 黄白线未同步新高 | 死叉减仓/清仓 |')
        w(f'| S2二卖 | 反弹不破一卖高点 | 一卖后反弹无力 | 反弹段MACD无力回0轴上方 | 再次死叉清仓 |')
        w(f'| S3三卖 | 跌破中枢下沿反抽 | 中枢下沿+反抽不涨回 | 跌破时MACD下0轴 | 跌破反抽离场 |')
        w()

        # ─── 摘要 ───
        w('## 复盘摘要')
        w()

        # 一、大盘速览
        w('### 大盘速览')
        w()
        regime_icon = {'牛市': '🟢', '震荡': '🟡', '熊市': '🔴'}.get(m.regime, '⚪')
        sentiment_icon = {'高潮': '🔥', '修复': '🌤', '分歧': '⚡', '冰点': '❄', '退潮': '🌧'}.get(m.sentiment_cycle, '➖')
        w(f'| 指数 | 涨跌 | 量能 | 市场状态 | 情绪 | 涨跌比 | 涨停/跌停 | 熊市风险 |')
        w(f'|------|------|------|----------|------|--------|-----------|----------|')
        w(f'| {m.index_close:.1f} | {m.index_change_pct:+.2f}% | {m.index_volume} | '
          f'{regime_icon} {m.regime} | {sentiment_icon} {m.sentiment_cycle} | '
          f'{m.up_pct:.0f}% | {m.limit_up_count}/{m.limit_down_count} | '
          f'{"⚠" if m.bear_risk else "✓"} |')
        w()

        # 二、结构全景
        w('### 结构全景')
        w()
        total = max(summary["total_stocks"], 1)
        w(f'| 趋势分布 | 买点信号 | 卖点信号 | 高评级 | 高危 |')
        w(f'|----------|----------|----------|--------|------|')
        w(f'| ↑{summary["up_trend"]}({summary["up_trend"]/total*100:.1f}%) '
          f'→{summary["consolidation"]}({summary["consolidation"]/total*100:.1f}%) '
          f'↓{summary["down_trend"]}({summary["down_trend"]/total*100:.1f}%) | '
          f'B1:{sum(1 for d in report.all_diagnoses if d.buy_point==1)} '
          f'B2:{summary["second_buy"]} '
          f'B3:{sum(1 for d in report.all_diagnoses if d.buy_point==3)} '
          f'(共{summary["buy_signals"]}) | '
          f'{summary["sell_signals"]}只 | '
          f'A:{summary["a_rank"]} B:{summary["b_rank"]} | '
          f'{summary["high_risk"]}只 |')
        w()

        # 三、主线行业
        w('### 主线行业')
        w()
        main_lines = [ind for ind in report.main_lines if ind.summary.startswith('★')][:5]
        sub_lines = [ind for ind in report.main_lines if ind.summary.startswith('○')][:3]
        avoid_lines = [ind for ind in report.main_lines if ind.summary.startswith('✗')]
        w(f'- **强主线**({len(main_lines)}个): {", ".join(ind.industry for ind in main_lines) if main_lines else "无"}')
        w(f'- **次主线**({len(sub_lines)}个): {", ".join(ind.industry for ind in sub_lines) if sub_lines else "无"}')
        if avoid_lines:
            w(f'- **回避方向**: {", ".join(ind.industry for ind in avoid_lines)}')
        w()

        # 四、操作建议汇总
        w('### 今日操作建议')
        w()
        # 仓位建议
        if m.sentiment_cycle in ('冰点', '退潮'):
            pos_advice = '防御仓位 (1-3成)，等待企稳信号'
        elif m.sentiment_cycle == '高潮':
            pos_advice = '积极仓位 (5-7成)，关注分歧信号'
        elif m.bear_risk:
            pos_advice = '轻仓或空仓 (0-2成)'
        else:
            pos_advice = '中性仓位 (3-5成)，主线内择优'
        w(f'- **仓位**: {pos_advice}')
        # 方向建议
        if main_lines:
            top_ind = main_lines[0]
            w(f'- **主攻**: {top_ind.industry}（B3{top_ind.b3_count}只+B1{top_ind.b1_count}只，梯队最完整）')
        # 回避
        sell_warning = summary['sell_signals'] > summary['buy_signals']
        if sell_warning:
            w(f'- **回避**: 卖点信号多于买点，非主线标的谨慎参与')
        w(f'- **关注**: 双级别确认买点{summary["confirmed_buy"]}只，优先关注')
        w()

        # 五、进化状态
        ah = report.accuracy_history
        if ah:
            recent = ah[-1]
            trend_icon = '✓' if recent['direction_accuracy'] >= 0.5 else '✗'
            w(f'- **准确率**: 最近 {trend_icon} {recent["direction_accuracy"]:.0%}，超额收益 {recent["excess_return"]:+.2f}%')
        if report.evolution_patches:
            urgent = [p for p in report.evolution_patches if p.urgency == 'urgent']
            if urgent:
                w(f'- **进化**: {len(urgent)}条紧急建议待处理')
        w()

        # 六、数据链接
        date_str = report.date.replace('-', '')
        w(f'> 📊 全市场个股明细: [缠论全量分析_{date_str}.csv](缠论全量分析_{date_str}.csv)')
        w(f'> 📈 准确率历史: accuracy_history.json')
        w()

        w('---')
        w('*报告由量化系统自动生成 | 缠论复盘引擎*')

        with open(path, 'w', encoding='utf-8') as f:
            f.write('\n'.join(lines))

    # ========== 系统自进化 ==========

    def _apply_evolution(self, patches: List[EvolutionPatch], date: str, auto_confirm: bool = False):
        """应用进化补丁 — 交互确认 + 自动更新配置文件"""
        # 加载进化历史
        history = []
        if os.path.exists(EVOLUTION_LOG):
            try:
                with open(EVOLUTION_LOG, 'r') as f:
                    history = json.load(f)
            except Exception:
                history = []

        # 仅处理 urgent/high 级别的补丁
        actionable = [p for p in patches if p.urgency in ('urgent', 'high')]
        if not actionable:
            return

        # 已在本轮或历史中应用过的补丁去重
        applied_keys = {h['key'] for h in history if h.get('applied')}
        new_patches = [p for p in actionable if p.key not in applied_keys]

        if not new_patches:
            print('\n  所有进化建议已在历史中应用，跳过。')
            return

        print(f'\n  ╔{"═"*74}╗')
        print(f'  ║  系统进化建议 ({len(new_patches)} 条){" " * 52}║')
        print(f'  ╠{"═"*74}╣')
        for p in new_patches:
            print(f'  ║  路径: {p.section}.{p.key}{" " * (64 - len(p.section) - len(p.key))}║')
            print(f'  ║  {p.current_value} → {p.suggested_value}  (置信度 {p.confidence:.0%}){" " * (46 - len(str(p.current_value)) - len(str(p.suggested_value)))}║')
            print(f'  ║  理由: {p.reason[:62]:<62s}║')
            print(f'  ║{"─"*74}║')

        # 交互确认或自动应用
        applied = []
        skipped = []
        for p in new_patches:
            applied_flag = False
            if auto_confirm:
                applied_flag = True
                print(f'\n  [自动确认] {p.key}: {p.current_value} → {p.suggested_value}')
            else:
                try:
                    ans = input(f'\n  应用 {p.key}: {p.current_value} → {p.suggested_value}? [y/N/q]: ').strip().lower()
                    if ans == 'q':
                        print('  已取消剩余补丁。')
                        break
                    applied_flag = ans == 'y'
                except (EOFError, KeyboardInterrupt):
                    print()
                    break

            if applied_flag:
                success = self._write_config_value(p.section, p.key, p.suggested_value)
                if success:
                    applied.append(p)
                else:
                    skipped.append(p)
            else:
                skipped.append(p)

        # 记录到进化历史
        for p in new_patches:
            applied_flag = p in applied
            history.append({
                'date': date,
                'section': p.section,
                'key': p.key,
                'from': p.current_value,
                'to': p.suggested_value,
                'reason': p.reason,
                'confidence': p.confidence,
                'urgency': p.urgency,
                'applied': applied_flag,
            })

        with open(EVOLUTION_LOG, 'w') as f:
            json.dump(history, f, ensure_ascii=False, indent=2)

        if applied:
            print(f'\n  ✓ 已应用 {len(applied)} 条: {", ".join(p.key for p in applied)}')
        if skipped:
            print(f'  ✗ 已跳过 {len(skipped)} 条: {", ".join(p.key for p in skipped)}')

    def _write_config_value(self, section: str, key: str, value) -> bool:
        """安全写入 YAML 配置文件（保留注释和格式）"""
        import yaml
        config_path = os.path.join(STRATEGY_DIR, 'config', 'factor_config.yaml')
        backup_path = config_path + '.backup'

        try:
            # 备份原文件
            import shutil
            shutil.copy2(config_path, backup_path)

            # 读取并修改
            with open(config_path, 'r', encoding='utf-8') as f:
                config = yaml.safe_load(f)

            # 导航到目标节点
            node = config
            parts = section.split('.')
            for part in parts:
                if part not in node:
                    node[part] = {}
                node = node[part]
            node[key] = value

            # 写回
            with open(config_path, 'w', encoding='utf-8') as f:
                yaml.dump(config, f, allow_unicode=True, default_flow_style=False, sort_keys=False)

            print(f'    已写入 {section}.{key} = {value} (备份: {backup_path})')
            return True
        except Exception as e:
            print(f'    写入失败: {e}')
            return False


# ==================== 入口 ====================

def main():
    parser = argparse.ArgumentParser(description='量化系统每日复盘 & 自我进化引擎')
    parser.add_argument('--date', type=str, default=None, help='复盘日期 (YYYY-MM-DD)')
    parser.add_argument('--top', type=int, default=15, help='显示 top N (默认15)')
    parser.add_argument('--symbols', type=str, default=None, help='指定股票，逗号分隔')
    parser.add_argument('--min-volume', type=float, default=None, help='最小成交额过滤')
    parser.add_argument('--evolve', action='store_true', default=False,
                        help='启用系统自进化（交互确认每条补丁）')
    parser.add_argument('--auto-evolve', action='store_true', default=False,
                        help='自动应用所有进化补丁（无需确认）')
    parser.add_argument('--auto-update', action='store_true', default=False,
                        help='自动更新数据+选股后再复盘（一条命令完成全流程）')
    parser.add_argument('--llm-report', action='store_true', default=False,
                        help='强制生成公众号版报告（API Key可用时默认自动生成）')
    parser.add_argument('--output', type=str, default=None, help='输出目录')

    args = parser.parse_args()

    symbols = None
    if args.symbols:
        symbols = [s.strip() for s in args.symbols.split(',')]

    engine = QuantReviewEngine(min_volume=args.min_volume)
    if args.output:
        engine.output_dir = args.output

    engine.run(date=args.date, top_n=args.top, symbols=symbols,
               evolve=args.evolve or args.auto_evolve, llm_report=args.llm_report,
               auto_evolve=args.auto_evolve, auto_update=args.auto_update)


if __name__ == '__main__':
    main()
